#!/usr/bin/env python3
"""Build the WSQ Learner Guide (LG) for Agile Project Management for Business.

Emits BOTH from one source so they can never diverge:
  * LG-<TITLE>.md            (Markdown mirror, at the repo root)
  * courseware/WSQ - Learner Guide - <TITLE> - <VER>.docx

The LG is the ONLY artifact carrying detailed step-by-step activity instructions
(the deck deliberately does not), plus the case brief, the datasets the activities
need, and the per-activity worksheets.
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
from prodoc import (add_cover_page, add_version_control, add_toc, add_page_numbers,
                    style_headings, enable_update_fields, _shade_cell, BRAND, DARK, GREY)

ACTIVITIES = DOMAIN1 + DOMAIN2 + DOMAIN3

# the real tool screenshot embedded in each activity's step-by-step section
TOOL_SHOT = {
    "Design Thinking": "tool-designthinking.png",
    "Fishbone": "tool-fishbone.png",
    "Scrum Board": "tool-scrum.png",
    "RACI Matrix": "tool-raci.png",
    "5 Whys": "tool-5whys.png",
    "Pareto Chart": "tool-paretochart.png",
}


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
SKILL_ASSETS = os.path.join(os.path.dirname(HERE), "assets")
CHARTS = os.path.join(REPO, "courseware", "assets")

MD = []          # markdown mirror lines


def md(line=""):
    MD.append(line)


doc = Document()
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.85)
    s.top_margin = s.bottom_margin = Inches(0.75)
style_headings(doc)

LOGO = os.path.join(SKILL_ASSETS, "tertiary-infotech-logo.png")
add_cover_page(doc, "Learner Guide", C.TITLE, C.VERSION, org_logo=LOGO,
               course_code=C.COURSE_CODE)
add_version_control(doc, [
    ("v9.0", C.VERSION_DATE,
     "Full content revamp. Expanded Agile theory from the Agile Manifesto, Scrum, Lean, "
     "Kanban, XP and DSDM. Added the five Agile metrics with worked interpretation and "
     "Little's Law. Replaced generic exercises with 8 detailed tool-based activities on one "
     "running case study (HarbourFront Logistics), each with step-by-step instructions, "
     "datasets, self-check criteria and a worksheet.",
     "Dr Alfred Ang"),
    ("v8.0", "28 November 2022", "Previous released version.", "Tertiary Infotech Academy"),
])
add_toc(doc)


# ------------------------------------------------------------------ helpers
def h1(t, mdlevel="##"):
    doc.add_heading(t, level=1)
    md(); md(f"{mdlevel} {t}"); md()


def h2(t):
    doc.add_heading(t, level=2)
    md(); md(f"### {t}"); md()


def h3(t):
    doc.add_heading(t, level=3)
    md(); md(f"#### {t}"); md()


def para(t, size=11, bold=False, italic=False, color=DARK, after=6, mdout=True):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    if mdout:
        if bold: md(f"**{t}**")
        elif italic: md(f"*{t}*")
        else: md(t)
        md()
    return p


def bullet(t, size=11, mdout=True):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.size = Pt(size)
    if mdout: md(f"- {t}")
    return p


def numbered(items, size=11):
    for i, t in enumerate(items, 1):
        p = doc.add_paragraph(style="List Number"); p.paragraph_format.space_after = Pt(4)
        r = p.add_run(t); r.font.size = Pt(size)
        md(f"{i}. {t}")
    md()


def callout(label, text, fill="EEF4FE", labelcolor=BRAND, size=11):
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]; cell.text = ""
    _shade_cell(cell, fill)
    p1 = cell.paragraphs[0]; p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(label.upper()); r1.bold = True; r1.font.size = Pt(9)
    r1.font.color.rgb = labelcolor
    p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(text); r2.font.size = Pt(size)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    md(); md(f"> **{label.upper()}** — {text}"); md()


def table(headers, rows, widths=None, header_fill="1F6FEB", font=9.5, mdout=True):
    t = doc.add_table(rows=0, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.add_row().cells
    for i, htext in enumerate(headers):
        hr[i].text = ""
        rr = hr[i].paragraphs[0].add_run(htext); rr.bold = True; rr.font.size = Pt(font)
        rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            for j, ln in enumerate(str(val).split("\n")):
                p = cells[i].paragraphs[0] if j == 0 else cells[i].add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                rr = p.add_run(ln); rr.font.size = Pt(font)
                if i == 0: rr.bold = True
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    if mdout:
        md(); md("| " + " | ".join(headers) + " |")
        md("|" + "|".join(["---"] * len(headers)) + "|")
        for row in rows:
            md("| " + " | ".join(str(c).replace("\n", "<br>") for c in row) + " |")
        md()
    return t


def image(name, width=6.2, caption=None):
    p = os.path.join(CHARTS, name)
    if os.path.exists(p):
        pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pp.add_run().add_picture(p, width=Inches(width))
        if caption:
            cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cr = cp.add_run(caption); cr.italic = True; cr.font.size = Pt(9)
            cr.font.color.rgb = GREY
        md(); md(f"![{caption or name}](courseware/assets/{name})")
        if caption: md(); md(f"*{caption}*")
        md()


def pagebreak():
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


# ================================================================== MD header
md(f"# {C.TITLE} — Learner Guide")
md()
md(f"**WSQ Course Code:** {C.COURSE_CODE}  ")
md(f"**Version:** {C.VERSION} · {C.VERSION_DATE}  ")
md(f"**Conducted by:** {C.ORG} ({C.UEN})  ")
md(f"**Duration:** {C.DAYS} days · {C.DURATION_HRS} instructional hours  ")
md(f"**TSC:** {C.TSC_TITLE} ({C.TSC_CODE})  ")
md()
md("---")

# ================================================================== 1. about
h1("1.  About This Course")
para("This Learner Guide is your reference during the course and after it. It carries the "
     "concepts, the running case study, the complete step-by-step instructions for all eight "
     "activities, the datasets those activities need, and a worksheet for each one. It is also "
     "one of the materials you may use in the open-book assessment.")

h2("1.1  Learning Outcomes")
para("By the end of this course you will be able to:")
for lo in C.LEARNING_OUTCOMES:
    bullet(lo, size=11)

h2("1.2  Course Structure")
table(["Topic", "Title", "Focus", "LO", "TSC", "Activities"],
      [[f"Topic {t['num']}", t["title"], t["subtitle"], t["lo"], t["tsc"],
        ", ".join(str(a["num"]) for a in ACTIVITIES if a["topic"] == t["num"])]
       for t in C.TOPICS],
      widths=[0.62, 1.55, 2.1, 0.42, 0.85, 0.72], font=9)

h2("1.3  How You Will Be Assessed")
table(["Instrument", "What it covers", "Format"], [
    ["Written Assessment (WA)", "The underpinning knowledge — the business landscape, "
     "customer analysis methods, Agile values and principles, Agile methodologies, change "
     "management, team formation models (K1–K7).",
     "Short-answer questions\n1 hour · open book"],
    ["Case Study (CS)", "Applied ability — reading a business scenario, ordering a backlog, "
     "interpreting metrics, diagnosing a root cause and recommending action (A1–A7).",
     "Applied business scenario\n1 hour · open book"],
], widths=[1.5, 3.9, 1.45], header_fill="F59E0B", font=9.5)
callout("Open book means", C.ASSESSMENT["open_book"] + " " + C.ASSESSMENT["note"],
        fill="FFF7E6", labelcolor=RGBColor(0xB4, 0x6A, 0x00))

h2("1.4  The Tools You Will Use")
para("All six tools run in a web browser. There is nothing to install and no account to "
     "create. Open each one when the activity calls for it.")
table(["Tool", "URL", "Used in"],
      [[t["name"], t["url"], t["use"]] for t in C.ED_TOOLS],
      widths=[1.25, 2.85, 2.75], header_fill="10B981", font=9.5)

pagebreak()

# ================================================================== 2. the case
h1("2.  The Running Case Study — HarbourFront Logistics")
para("Every activity in this course works on the same business case, so what you build in one "
     "activity is the input to the next. Read this brief once, carefully — you will return to "
     "it eight times, and it is the same style of scenario used in the Case Study assessment.")

h2("2.1  The Company")
para("HarbourFront Logistics Pte Ltd is a Singapore third-party logistics (3PL) provider with "
     "140 staff. It handles inbound sea and air freight for manufacturing customers, "
     "manages customs clearance, and runs a bonded warehouse in Jurong. Its customers are "
     "mostly mid-sized manufacturers who depend on predictable inbound flow to keep their "
     "production lines running.")

h2("2.2  What Happened to CustomerConnect")
para("In January 2024 HarbourFront approved 'CustomerConnect', a self-service portal intended "
     "to let customers track their own shipments and cut inbound calls to the service desk. "
     "The project ran as a conventional waterfall delivery:")
table(["Fact", "Detail"], [
    ["Approach", "Waterfall — requirements, design, build, test, deploy"],
    ["Requirements", "A 96-page requirements document, signed off in month 1 and frozen"],
    ["Planned duration", "11 months"],
    ["Actual duration", "15 months — delivered 4 months late"],
    ["Change requests", "41 raised during the build; 41 rejected as out of scope"],
    ["Customer contact", "Requirements workshops in month 1, then nothing until UAT in month 13"],
    ["First integration", "Month 9 — the components had never run together before that point"],
    ["Testing", "Compressed into the final 6 weeks"],
    ["Features delivered", "24"],
    ["Features used weekly, 6 months after launch", "3"],
    ["Top complaint to the service desk", "“Where is my shipment right now?”"],
    ["Inbound call volume", "Unchanged from before the portal launched"],
], widths=[2.35, 4.45], header_fill="DC2626", font=9.5)

callout("The uncomfortable part", "The project delivered exactly what was signed off. Every "
        "requirement in the 96-page document was built and accepted at UAT. It still failed, "
        "because the document described what people could imagine in month 1 — not what "
        "customers needed in month 15.", fill="FDECEC",
        labelcolor=RGBColor(0x9B, 0x2C, 0x2C))

h2("2.3  Where You Come In")
para("You have just been appointed Agile project lead for the CustomerConnect restart. The "
     "management committee has approved 6 sprints of 2 weeks, one team of 7 people, and a "
     "demonstration to the top 5 customers in 12 weeks. Nothing from the old requirements "
     "document carries over unexamined.")

h2("2.4  The People")
table(["Person", "Role", "What you need to know"], [
    ["Priya Menon", "Warehouse operations executive at a customer",
     "Tracks 40–60 inbound shipments a week. Her production line stops if she mis-times a "
     "delivery. Your primary user."],
    ["Lim Wei Sheng", "HarbourFront operations manager",
     "Acts as customer proxy internally. Well-intentioned, but gives the team priorities "
     "directly — which causes the conflict you resolve in Activity 4."],
    ["Rachel Tan", "Programme director (legacy)",
     "Ran the original waterfall project. Still approving story changes and asking for "
     "percent-complete reports."],
    ["The delivery team", "7 people",
     "Two developers, one tester, one designer, one business analyst, one data engineer, "
     "one Scrum Master. One specialist per skill — which creates the queue you find in "
     "Activity 8."],
], widths=[1.15, 1.55, 4.1], font=9.5)

pagebreak()

# ================================================================== 3..5 topics
TOPIC_INTRO = {
    1: ("Topic 1 establishes why Agile exists. You will look honestly at the business "
        "operating landscape, understand what the waterfall model assumes and where those "
        "assumptions break, and see what Agile actually replaces them with. You will also "
        "learn when Agile is the wrong answer — an Agile practitioner who cannot say that "
        "is selling, not advising."),
    2: ("Topic 2 is the essentials — the Agile Manifesto's four values and twelve principles, "
        "then the frameworks built on them: Scrum in depth, Lean, Kanban and XP. It closes "
        "with the part most courses skip: how to find support for Agile inside a real "
        "organisation, and how to handle the resistance you will certainly meet."),
    3: ("Topic 3 is where Agile becomes a week of actual work. You will build the team, set "
        "the vision, write and estimate stories, run a sprint, and read the five metrics that "
        "tell you what is really happening. Most importantly, you will learn to diagnose a "
        "delivery problem to a cause you can change."),
}

for t in C.TOPICS:
    h1(f"{2 + t['num']}.  Topic {t['num']} — {t['title']}")
    para(TOPIC_INTRO[t["num"]])
    table(["Learning outcome", "TSC coverage", "Activities in this topic"],
          [[t["lo"], t["tsc"], ", ".join(f"Activity {a['num']}" for a in ACTIVITIES
                                         if a["topic"] == t["num"])]],
          widths=[1.5, 2.4, 2.9], header_fill="7C3AED", font=9.5)

    h2(f"{2 + t['num']}.1  Key Concepts")
    for i, cpt in enumerate(t["concepts"], 1):
        bullet(cpt, size=11)

    # topic-specific figures + reference tables
    if t["num"] == 1:
        h2("3.2  Figures for Topic 1")
        image("waterfall-vs-agile.png", 6.3,
              "Waterfall concentrates risk and value at the end; Agile releases both continuously.")
        image("cost-of-change.png", 5.6,
              "The cost-of-change curve — the economic argument for short iterations.")
        image("iron-triangle.png", 5.6,
              "Inverting the triangle: Agile fixes time and cost, and varies scope.")
        image("value-delivery.png", 5.6,
              "Cumulative value delivered: Agile has released ~34% of value by sprint 4.")
        h2("3.3  Reference — Agile vs Waterfall")
        table(["Aspect", "Waterfall / Traditional", "Agile"],
              [["Approach", "Sequential phases, one pass", "Iterative and incremental"],
               ["Requirements", "Fixed and signed off upfront", "Emerge and are re-ordered continuously"],
               ["Flexibility", "Resists change after planning", "Adapts throughout"],
               ["Customer role", "Consulted at the start and at UAT", "Continuous collaboration"],
               ["Delivery", "One release at the end", "A usable increment every sprint"],
               ["Risk", "Addressed upfront, realised late", "Retired continuously"],
               ["Documentation", "Comprehensive by mandate", "Barely sufficient, just in time"],
               ["Team structure", "Hierarchical, specialised roles", "Self-organising, cross-functional"],
               ["Best suited to", "Stable scope, strict compliance, known technology",
                "Evolving requirements, high novelty, need for fast feedback"]],
              widths=[1.35, 2.65, 2.8], font=9.5)
        h2("3.4  Reference — Choosing Your Approach")
        table(["If this is true of your project…", "Lean towards"], [
            ["Requirements are genuinely fixed by regulation or contract", "Waterfall"],
            ["The customer cannot give feedback more than once", "Waterfall"],
            ["An increment cannot be built and changed cheaply (civil works, hardware tooling)",
             "Waterfall or staged hybrid"],
            ["Requirements are uncertain or will evolve", "Agile"],
            ["The work is novel and nobody has built it before", "Agile"],
            ["Early partial value is worth more than complete late value", "Agile"],
            ["You need to reduce the risk of building the wrong thing", "Agile"],
            ["Governance can fund incrementally and stop early", "Agile"],
            ["Some of the above, but not all", "A deliberate hybrid — and say which parts are which"],
        ], widths=[4.6, 2.2], header_fill="10B981", font=9.5)

    elif t["num"] == 2:
        h2("4.2  Figures for Topic 2")
        image("manifesto-values.png", 6.3, "The four values of the Agile Manifesto (2001).")
        image("scrum-framework.png", 6.3,
              "The Scrum framework: 3 accountabilities, 3 artefacts, 5 events, one feedback loop.")
        image("definition-of-done.png", 6.0,
              "A Definition of Done — the team's shared, visible and testable quality bar.")
        image("lean-wastes.png", 6.2, "The eight wastes that Lean targets.")
        image("kanban-littles-law.png", 6.2,
              "A Kanban board with WIP limits, and Little's Law — the reason WIP limits work.")
        h2("4.3  Reference — The Twelve Principles of the Agile Manifesto")
        table(["#", "Principle"], [
            ["1", "Our highest priority is to satisfy the customer through early and continuous "
                  "delivery of valuable software."],
            ["2", "Welcome changing requirements, even late in development. Agile processes "
                  "harness change for the customer's competitive advantage."],
            ["3", "Deliver working software frequently, from a couple of weeks to a couple of "
                  "months, with a preference to the shorter timescale."],
            ["4", "Business people and developers must work together daily throughout the project."],
            ["5", "Build projects around motivated individuals. Give them the environment and "
                  "support they need, and trust them to get the job done."],
            ["6", "The most efficient and effective method of conveying information to and within "
                  "a development team is face-to-face conversation."],
            ["7", "Working software is the primary measure of progress."],
            ["8", "Agile processes promote sustainable development. The sponsors, developers and "
                  "users should be able to maintain a constant pace indefinitely."],
            ["9", "Continuous attention to technical excellence and good design enhances agility."],
            ["10", "Simplicity — the art of maximizing the amount of work not done — is essential."],
            ["11", "The best architectures, requirements and designs emerge from self-organizing teams."],
            ["12", "At regular intervals, the team reflects on how to become more effective, then "
                   "tunes and adjusts its behaviour accordingly."],
        ], widths=[0.4, 6.4], font=9.5)
        h2("4.4  Reference — Scrum at a Glance")
        table(["Element", "Who / what", "Purpose"], [
            ["Product Owner", "One person", "Owns value and the ORDER of the product backlog. "
             "Accepts or rejects each increment."],
            ["Scrum Master", "One person", "Owns the process. Servant leader; removes impediments; "
             "coaches the team."],
            ["Developers", "Typically 3–9 people", "Own HOW the work is done and how much enters "
             "the sprint."],
            ["Product Backlog", "Artefact", "The ordered list of everything known to be needed. "
             "Commitment: the Product Goal."],
            ["Sprint Backlog", "Artefact", "The selected items plus the plan. Commitment: the "
             "Sprint Goal."],
            ["Increment", "Artefact", "A usable, integrated slice. Commitment: the Definition of Done."],
            ["Sprint", "Event (container)", "Fixed length, 1–4 weeks. Does not change length "
             "between sprints."],
            ["Sprint Planning", "Event, up to 8h for a 1-month sprint", "Decide the Sprint Goal, "
             "what is selected and how it will be built."],
            ["Daily Scrum", "Event, 15 minutes", "The Developers re-plan the next 24 hours."],
            ["Sprint Review", "Event", "Inspect the increment with stakeholders and adapt the backlog."],
            ["Sprint Retrospective", "Event", "Inspect the process and commit to one or two "
             "improvements."],
        ], widths=[1.3, 1.6, 3.9], header_fill="7C3AED", font=9.5)
        h2("4.5  Reference — Scrum vs Kanban vs Scrumban")
        table(["Dimension", "Scrum", "Kanban", "Scrumban"], [
            ["Cadence", "Fixed sprints, 1–4 weeks", "Continuous flow", "Cadence plus WIP limits"],
            ["Commitment", "A Sprint Goal per sprint", "Pull when capacity frees", "Goal, flexible pull"],
            ["Best for", "Teams of 5–9 with a product goal", "Support queues, high variability",
             "Teams outgrowing strict Scrum"],
            ["Roles", "PO, SM, Developers", "None prescribed", "Usually keeps PO and SM"],
            ["Key metrics", "Velocity, sprint burndown", "Cycle time, throughput, CFD", "Both sets"],
            ["Change mid-cycle", "Protected by the Sprint Goal", "Allowed any time within WIP limits",
             "Allowed within WIP limits"],
        ], widths=[1.15, 1.9, 1.9, 1.85], header_fill="10B981", font=9)

    else:
        h2("5.2  Figures for Topic 3")
        image("tuckman-leadership.png", 6.3,
              "Tuckman's stages, with the leadership style each stage needs.")
        image("user-story-anatomy.png", 6.2,
              "The anatomy of a user story, its acceptance criteria and the INVEST test.")
        image("sprint-burndown.png", 5.9,
              "Sprint burndown — the flat section and the carryover are the informative parts.")
        image("velocity-forecast.png", 6.3,
              "Velocity, and a release forecast expressed as a range rather than a date.")
        image("control-chart.png", 6.0,
              "Control chart — rising velocity together with rising cycle time means rising WIP.")
        image("cumulative-flow.png", 6.0,
              "Cumulative flow diagram — the widening band identifies the bottleneck.")
        image("retro-stages.png", 6.3, "The retrospective in five stages.")
        image("pareto-defects.png", 6.2,
              "Pareto analysis of the CustomerConnect defect data used in Activity 7.")
        image("value-stream-map.png", 6.3,
              "A value stream map — 13 days of work inside 125 days of lead time.")
        h2("5.3  Reference — The Five Agile Metrics")
        table(["Metric", "What it shows", "How to read it", "The anti-pattern"], [
            ["Sprint burndown", "Points remaining in one sprint, day by day",
             "A flat line means work is started but not finished — check WIP, not effort",
             "Treating a late vertical drop as success; it means late integration"],
            ["Epic / release burndown", "Progress across a release or epic",
             "Slope gives the completion trend; scope added shows as the line rising",
             "Hiding scope growth by rebaselining silently"],
            ["Velocity", "Average points DONE per sprint",
             "Use the slowest, average and fastest sprints to build a forecast range",
             "Using it as a productivity target — this causes estimate inflation"],
            ["Control chart", "Cycle time and lead time per item, and the trend",
             "Rising cycle time with rising velocity means WIP is rising",
             "Celebrating velocity while cycle time worsens"],
            ["Cumulative flow diagram", "Count of items per status over time",
             "A band widening vertically is a bottleneck at that status",
             "Adding capacity to the busiest-looking column instead of the constraint"],
        ], widths=[1.15, 1.5, 2.2, 1.95], header_fill="1F6FEB", font=8.5)
        h2("5.4  Reference — Little's Law Worked")
        table(["Scenario", "WIP", "Throughput", "Cycle time"], [
            ["The team starts everything", "30 items", "5 items/week", "6 weeks"],
            ["The team halves WIP", "15 items", "5 items/week", "3 weeks"],
            ["The team quarters WIP", "8 items", "5 items/week", "1.6 weeks"],
        ], widths=[2.6, 1.4, 1.5, 1.3], header_fill="10B981", font=9.5)
        callout("Read that table again", "Throughput is identical in all three rows. Only WIP "
                "changed, and cycle time fell with it. Cycle Time = WIP ÷ Throughput is "
                "arithmetic, not a management opinion — which is why limiting WIP is the single "
                "most reliable way to deliver faster.")
        h2("5.5  Reference — Root-Cause Tools, and When to Use Each")
        table(["Tool", "Use it when", "It gives you"], [
            ["Fishbone (Ishikawa)", "You need to see the WHOLE problem space across categories",
             "A structured spread of candidate causes — breadth"],
            ["Pareto chart", "You have many causes and limited capacity",
             "Which few causes carry most of the pain — priority"],
            ["5 Whys", "You have chosen ONE problem and need the cause you can change",
             "A single causal chain to a changeable root cause — depth"],
            ["PDCA / Kaizen", "You have a root cause and need it fixed",
             "A funded, owned improvement carried into the next sprint"],
        ], widths=[1.25, 2.9, 2.65], header_fill="7C3AED", font=9.5)
        callout("The sequence that works", "Fishbone to see the whole space → Pareto to pick the "
                "few that matter → 5 Whys to reach the changeable cause → one SMART action with "
                "an owner and story points in the next sprint backlog. You will do exactly this "
                "in Activities 2, 6 and 7.")

    pagebreak()

# ================================================================== 6. activities
h1("6.  Activities — Step-by-Step Instructions")
para("This section carries the complete instructions for all eight activities. Work in your "
     "team, follow the numbered steps in order, and use the self-check at the end of each "
     "activity before you move on. Each activity also has its own folder in the course "
     "materials containing a worksheet and any data files you need.")
callout("Why the slides do not carry these steps",
        "The trainer's slides show the situation, the tool and the expected outcome. The "
        "detailed procedure lives here, in your Learner Guide, so you can work at your own "
        "pace during the activity and refer back to it after the course.")

table(["#", "Activity", "Tool", "Min", "Topic", "LO"],
      [[str(a["num"]), a["title"], a["tool"], str(a["duration"]),
        f"Topic {a['topic']}", a["lo"]] for a in ACTIVITIES],
      widths=[0.32, 2.85, 1.15, 0.42, 0.72, 0.42], header_fill="10B981", font=9.5)

pagebreak()

for a in ACTIVITIES:
    h2(f"6.{a['num']}  Activity {a['num']} — {a['title']}")
    table(["Field", "Detail"], [
        ["Topic", f"Topic {a['topic']} — {[t['title'] for t in C.TOPICS if t['num'] == a['topic']][0]}"],
        ["Learning outcome", a["lo"]],
        ["Objective", a["objective"]],
        ["Duration", f"{a['duration']} minutes"],
        ["Tool", f"{a['tool']} — {a['tool_url']}"],
        ["Team size", "3–4 learners"],
        ["Materials", a["services"]],
        ["Activity folder", f"activities/activity-{a['num']:02d}-*/"],
    ], widths=[1.35, 5.45], font=9.5)

    h3("The situation")
    para(a["scenario"], size=11)

    h3("What you will do")
    para(a["desc"], size=11)

    h3("What you will produce")
    para(a["build"], size=11, bold=True)

    h3(f"The tool — {a['tool']}")
    para(f"Open {a['tool_url']} in your browser. This is the screen you will see when the "
         f"tool loads; follow the numbered steps below from here.", size=11)
    image(TOOL_SHOT[a["tool"]], 6.1,
          caption=f"{a['tool']} — {a['tool_url']}")

    h3("Step-by-step instructions")
    numbered([s[0] for s in a["steps"]], size=11)

    h3("Self-check — are you done?")
    callout("Done when", a["test"], fill="E8F7EE",
            labelcolor=RGBColor(0x12, 0x7A, 0x3E))

    h3("Debrief — what this activity proves")
    for d in a["debrief"]:
        bullet(d, size=11)

    # activity-specific datasets
    if a["num"] == 7:
        h3("Dataset — CustomerConnect defect causes, Sprints 1–3")
        para("Enter these 8 categories into the Pareto tool. The total must come to 120.",
             size=11)
        table(["Defect cause category", "Count"], [
            ["Stale or missing feed data", "44"],
            ["Unclear acceptance criteria", "31"],
            ["Environment and configuration drift", "17"],
            ["Carrier API contract changes", "11"],
            ["UI validation gaps", "8"],
            ["Access and permission errors", "5"],
            ["Report formatting", "3"],
            ["Documentation errors", "1"],
            ["TOTAL", "120"],
        ], widths=[4.4, 1.0], header_fill="DC2626", font=9.5)
        table(["Cumulative check", "Cumulative %"], [
            ["Stale/missing feed data", "36.7%"],
            ["+ Unclear acceptance criteria", "62.5%"],
            ["+ Environment drift", "76.7%"],
            ["+ Carrier API changes", "85.8%"],
        ], widths=[3.2, 1.4], header_fill="7C3AED", font=9.5)

    if a["num"] == 8:
        h3("Dataset — velocity and remaining backlog")
        table(["Input", "Value"], [
            ["Sprint 1 velocity", "14 points"],
            ["Sprint 2 velocity", "18 points"],
            ["Sprint 3 velocity", "17 points"],
            ["Sprint 4 velocity", "21 points"],
            ["Average velocity", "17.5 points"],
            ["Points remaining in the backlog", "186 points"],
            ["Sprint length", "2 weeks"],
        ], widths=[3.0, 1.6], header_fill="1F6FEB", font=9.5)
        h3("Dataset — cycle time by sprint (for the control chart)")
        table(["Sprint", "Average cycle time"], [
            ["Sprint 1", "3.1 days"], ["Sprint 2", "4.3 days"],
            ["Sprint 3", "5.8 days"], ["Sprint 4", "7.4 days"],
        ], widths=[1.6, 2.0], header_fill="7C3AED", font=9.5)
        h3("Expected answer — check your working")
        table(["Calculation", "Result"], [
            ["Average forecast: 186 ÷ 17.5", "10.6 → 11 sprints (~22 weeks)"],
            ["Pessimistic: 186 ÷ 14", "13.3 → 14 sprints (~28 weeks)"],
            ["Optimistic: 186 ÷ 21", "8.9 → 9 sprints (~18 weeks)"],
            ["Answer to the committee", "18–28 weeks, most likely ~22 weeks, "
                                        "assuming a stable team and no backlog growth"],
            ["Bottleneck from the CFD", "Testing — its band widens vertically over time"],
            ["Confirmation from the control chart", "Cycle time rose 3.1 → 7.4 days while "
                                                    "velocity rose"],
            ["Cause", "WIP rose faster than throughput past the Testing constraint "
                      "(Little's Law)"],
            ["Recommended action", "Lower the WIP limit on Testing; stop starting new stories "
                                   "until Testing clears"],
        ], widths=[2.55, 4.25], header_fill="10B981", font=9)

    h3("Worksheet")
    para(f"Record your team's output for Activity {a['num']} below, or in the worksheet file in "
         f"the activity folder.", size=10, italic=True)
    ws = doc.add_table(rows=1, cols=1); ws.style = "Table Grid"
    wc = ws.rows[0].cells[0]
    wc.text = ""
    for _ in range(9):
        wp = wc.add_paragraph(); wp.paragraph_format.space_after = Pt(9)
        wr = wp.add_run("​"); wr.font.size = Pt(11)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    md(); md("**Worksheet:** record your team's output in the activity folder worksheet.")
    md()
    pagebreak()

# ================================================================== 7. glossary
h1("7.  Glossary")
GLOSSARY = [
    ("Acceptance criteria", "The conditions a user story must satisfy to be accepted, usually "
     "written Given / When / Then."),
    ("Agile Manifesto", "The 2001 statement of 4 values and 12 principles from which the Agile "
     "frameworks derive. agilemanifesto.org"),
    ("Burndown chart", "A chart of work remaining over time, within a sprint or across a release."),
    ("Cadence", "The fixed rhythm of a team's events — for example a 2-week sprint."),
    ("Control chart", "A chart of cycle time or lead time per work item, showing the trend."),
    ("Cumulative flow diagram (CFD)", "A stacked area chart of item counts per status over time; "
     "a widening band identifies a bottleneck."),
    ("Cycle time", "How long an item takes to move through the team's active process."),
    ("Daily Scrum", "A 15-minute daily event in which the Developers re-plan the next 24 hours."),
    ("Definition of Done (DoD)", "The team's shared, testable standard for calling work complete. "
     "Applies to every item."),
    ("DSDM / AgilePM", "The Agile framework endorsed by APM (UK), with feasibility, foundations, "
     "evolutionary development and deployment phases."),
    ("Epic", "A large body of work that is broken down into multiple user stories."),
    ("Impediment", "Anything blocking the team's progress that the team cannot remove itself; "
     "the Scrum Master owns its removal."),
    ("Increment", "A usable, integrated slice of product that meets the Definition of Done."),
    ("INVEST", "A test for a good user story: Independent, Negotiable, Valuable, Estimatable, "
     "Small, Testable."),
    ("Kaizen", "Continuous improvement through small, frequent, team-owned changes."),
    ("Kanban", "A flow-based method: visualise the workflow, limit WIP, manage flow, make "
     "policies explicit, improve collaboratively."),
    ("Lead time", "The customer's total wait, from request to delivery. Includes cycle time "
     "plus all queueing."),
    ("Lean", "A Toyota-derived approach focused on eliminating waste and maximising value."),
    ("Little's Law", "Cycle Time = Work in Progress ÷ Throughput. The reason limiting WIP "
     "speeds delivery."),
    ("MoSCoW", "A prioritisation method: Must have, Should have, Could have, Won't have this time."),
    ("MVP (Minimum Viable Product)", "The smallest release that delivers real value and produces "
     "genuine learning."),
    ("PDCA", "Plan-Do-Check-Act — Deming's improvement cycle, used to carry a retrospective "
     "action to completion."),
    ("Planning poker", "Relative estimation in which the team reveals Fibonacci estimates "
     "simultaneously to avoid anchoring."),
    ("Product Backlog", "The single ordered list of everything known to be needed in the product."),
    ("Product Owner", "The one person accountable for the value of the product and the order of "
     "the Product Backlog."),
    ("Refinement (grooming)", "The ongoing activity of adding detail, estimates and order to "
     "backlog items."),
    ("Retrospective", "The event at the end of a sprint in which the team inspects its process "
     "and commits to improvement."),
    ("Scrum", "The most widely used Agile framework: 3 accountabilities, 3 artefacts, 5 events, "
     "built on empiricism."),
    ("Scrum Master", "The person accountable for the team's process, coaching and impediment "
     "removal. A servant leader."),
    ("Scrumban", "A hybrid that keeps a Scrum cadence while adding Kanban WIP limits."),
    ("Servant leadership", "Leading by enabling — removing obstacles and providing what the team "
     "needs, rather than directing."),
    ("Spike", "A timeboxed investigation used to reduce a specific technical or risk unknown."),
    ("Sprint", "A fixed-length container of 1–4 weeks in which all Scrum events occur."),
    ("Sprint Backlog", "The Sprint Goal, the selected items, and the plan for delivering them."),
    ("Sprint Goal", "The single objective for a sprint. The commitment carried by the Sprint Backlog."),
    ("Story point", "A relative unit combining complexity, effort and risk."),
    ("Technical debt", "Work deferred to move faster now, which makes later work slower until repaid."),
    ("Throughput", "The number of items completed per unit of time."),
    ("Timebox", "A fixed maximum duration for an activity. Work is adjusted to fit the box."),
    ("TSC", "Technical Skills and Competencies — the Singapore Skills Framework unit this course "
     f"maps to ({C.TSC_CODE})."),
    ("User story", "A requirement expressed from the user's perspective: As a <role>, I want "
     "<goal>, so that <benefit>."),
    ("Velocity", "The average story points completed per sprint. A forecasting input, never a target."),
    ("VUCA", "Volatility, uncertainty, complexity and ambiguity — the conditions Agile is "
     "designed for."),
    ("Waterfall", "A sequential delivery approach in which each phase completes before the next "
     "begins."),
    ("WIP (Work in Progress)", "Work started but not finished. High WIP inflates cycle time and "
     "hides bottlenecks."),
    ("XP (Extreme Programming)", "An Agile method contributing TDD, pair programming, continuous "
     "integration and refactoring."),
]
table(["Term", "Meaning"], [[g[0], g[1]] for g in GLOSSARY],
      widths=[1.75, 5.05], font=9.5)

# ================================================================== 8. further reading
h1("8.  Further Reading and Sources")
para("The content of this guide draws on the following public sources, in addition to the "
     "course's own material:")
SOURCES = [
    ("Agile Manifesto", "https://agilemanifesto.org/",
     "The original four values and twelve principles."),
    ("Atlassian — Agile Project Management",
     "https://www.atlassian.com/agile/project-management",
     "Practical guidance on Agile delivery, and the five agile metrics used in Topic 3."),
    ("APM (Association for Project Management) — Agile Project Management",
     "https://www.apm.org.uk/resources/find-a-resource/agile-project-management/",
     "The UK professional body's definition, benefits, principles and governance guidance."),
    ("Coursera — What Is Agile? A Beginner's Guide",
     "https://www.coursera.org/articles/what-is-agile-a-beginners-guide",
     "The Agile lifecycle, methodology comparison and framework adoption figures."),
    ("Rasmussen — What Is Agile Project Management?",
     "https://www.rasmussen.edu/degrees/business/blog/what-is-agile-project-management/",
     "Agile in non-software business contexts, with company examples."),
    ("Adobe Business — Agile methodology: frameworks and best practices",
     "https://business.adobe.com/blog/basics/agile",
     "Agile applied to marketing and creative teams; scaling and team-size guidance."),
    ("GeeksforGeeks — Agile Project Management",
     "https://www.geeksforgeeks.org/software-engineering/agile-project-management/",
     "The five-phase APM lifecycle, advantages, disadvantages and comparison tables."),
]
table(["Source", "Link", "What it contributes"], SOURCES,
      widths=[1.85, 2.35, 2.6], font=9)

h2("8.1  Recommended Next Courses")
for r in C.RECOMMENDED:
    bullet(r, size=11)

h2("8.2  Support")
table(["Channel", "Detail"], [
    ["Email", "enquiry@tertiaryinfotech.com"],
    ["Telephone", "+65 6100 0613"],
    ["Website", "www.tertiarycourses.com.sg"],
    ["LMS / TMS", "https://lms-tms.tertiaryinfotech.com"],
], widths=[1.5, 5.3], font=10)

add_page_numbers(doc)
enable_update_fields(doc)

OUTDIR = os.path.join(REPO, "courseware")
os.makedirs(OUTDIR, exist_ok=True)
OUT = os.path.join(OUTDIR, f"WSQ - Learner Guide - {C.SHORT_TITLE} - {C.VERSION}.docx")
doc.save(OUT)

MDOUT = os.path.join(REPO, f"LG-{C.SHORT_TITLE}.md")
with open(MDOUT, "w") as f:
    f.write("\n".join(MD) + "\n")

print("Saved:", OUT)
print("Saved:", MDOUT)
print(f"Activities documented: {len(ACTIVITIES)}")
print(f"Total documented steps: {sum(len(a['steps']) for a in ACTIVITIES)}")
