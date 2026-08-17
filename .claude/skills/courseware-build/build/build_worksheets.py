#!/usr/bin/env python3
"""Build a printable PDF worksheet for every activity, from the same single source.

Each activity folder gets:

    activities/activity-NN-<slug>/
        WORKSHEET.md     the Markdown worksheet (already generated)
        WORKSHEET.docx   the same worksheet, house-styled and printable
        WORKSHEET.pdf    rendered from the DOCX

The Markdown fenced blocks (```…```) are placeholders a learner cannot write in.
In the DOCX/PDF each becomes a REAL ruled answer box — a bordered table cell with
writing lines — so the worksheet can be printed and filled in by hand, or typed
into on screen.

A combined "All Activity Worksheets" PDF is also produced for easy printing of
the whole set.
"""
import os
import re
import subprocess
import sys
import glob

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
from prodoc import add_page_numbers, style_headings, _shade_cell, BRAND, DARK, GREY

ACTIVITIES = DOMAIN1 + DOMAIN2 + DOMAIN3

TEAL = RGBColor(0x10, 0xB9, 0x81)
AMBER = RGBColor(0xF5, 0x9E, 0x0B)
VIOLET = RGBColor(0x7C, 0x3A, 0xED)

SOFFICE = "/Applications/LibreOffice.app/Contents/MacOS/soffice"


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
ACTDIR = os.path.join(REPO, "activities")
SKILL_ASSETS = os.path.join(os.path.dirname(HERE), "assets")
LOGO = os.path.join(SKILL_ASSETS, "tertiary-infotech-logo.png")


# ------------------------------------------------------------------ md parsing
def parse_worksheet(md_text):
    """Parse a generated WORKSHEET.md into its parts.

    Returns dict(title, meta[list of (label, value)], fields[list of (n, label)],
                 checks[list of str]).
    The Markdown is the single source — the PDF never invents content.
    """
    title = ""
    meta, fields, checks = [], [], []
    lines = md_text.splitlines()
    for i, ln in enumerate(lines):
        if ln.startswith("# ") and not title:
            title = ln[2:].strip()
        m = re.match(r"^\*\*(.+?):\*\*\s*(.*?)\s*$", ln)
        if m and not fields:
            meta.append((m.group(1).strip(), m.group(2).strip().rstrip("_ ").strip()))
        m = re.match(r"^###\s+(\d+)\.\s+(.*)$", ln)
        if m:
            fields.append((int(m.group(1)), m.group(2).strip()))
        m = re.match(r"^-\s*\[\s*\]\s*(.+)$", ln)
        if m:
            checks.append(m.group(1).strip())
    return dict(title=title, meta=meta, fields=fields, checks=checks)


# ------------------------------------------------------------------ docx helpers
def new_doc():
    d = Document()
    st = d.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(0.7)
        s.top_margin = Inches(0.6); s.bottom_margin = Inches(0.6)
    style_headings(d)
    return d


def line(d, text="", bold=False, size=11, color=DARK, after=6, align=None, italic=False):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    return p


def _set_border(cell, color="C9D3DF", sz="6"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single"); el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0"); el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def answer_box(d, rules=3, width=7.1):
    """A real bordered answer box containing ruled writing lines."""
    t = d.add_table(rows=1, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.rows[0].cells[0]
    cell.width = Inches(width)
    cell.text = ""
    _set_border(cell)
    for i in range(rules):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(7)
        pPr = p._p.get_or_add_pPr()
        bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1"); bottom.set(qn("w:color"), "D7E0EA")
        bdr.append(bottom); pPr.append(bdr)
        p.add_run("").font.size = Pt(11)
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def checkbox_line(d, text):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.left_indent = Inches(0.1)
    r = p.add_run("☐   "); r.font.size = Pt(13); r.font.color.rgb = TEAL; r.bold = True
    r2 = p.add_run(text); r2.font.size = Pt(10)
    return p


def header_block(d, a, ws):
    """Worksheet header: activity identity + team/date/tool fill-ins."""
    if os.path.exists(LOGO):
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(0)
        p.add_run().add_picture(LOGO, height=Inches(0.34))

    line(d, f"ACTIVITY {a['num']} WORKSHEET", bold=True, size=11, color=TEAL, after=2)
    line(d, a["title"], bold=True, size=16, color=DARK, after=3)
    line(d, f"{C.TITLE}  ·  {C.COURSE_CODE}  ·  Topic {a['topic']}  ·  {a['lo']}  ·  "
            f"{a['duration']} minutes",
         size=9.5, color=GREY, after=8)

    # team / date / tool
    t = d.add_table(rows=2, cols=2); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    rows = [("Team members", ""), ("Date", ""),
            ("Tool", a["tool_url"]), ("Activity folder", f"activities/activity-{a['num']:02d}-*/")]
    cells = [t.rows[0].cells[0], t.rows[0].cells[1], t.rows[1].cells[0], t.rows[1].cells[1]]
    for cell, (label, val) in zip(cells, rows):
        cell.text = ""
        _shade_cell(cell, "F5F8FC")
        p1 = cell.paragraphs[0]; p1.paragraph_format.space_after = Pt(1)
        r1 = p1.add_run(label.upper()); r1.bold = True; r1.font.size = Pt(8)
        r1.font.color.rgb = BRAND
        p2 = cell.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
        if val:
            r2 = p2.add_run(val); r2.font.size = Pt(9.5)
        else:
            r2 = p2.add_run("____________________________________"); r2.font.size = Pt(10)
            r2.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    d.add_paragraph().paragraph_format.space_after = Pt(4)

    # the situation, so the worksheet stands alone away from the guide
    t2 = d.add_table(rows=1, cols=1); t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = t2.rows[0].cells[0]; c.text = ""
    _shade_cell(c, "F5F8FC")
    p = c.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run("THE SITUATION"); r.bold = True; r.font.size = Pt(8.5)
    r.font.color.rgb = VIOLET
    p2 = c.add_paragraph(); p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(a["scenario"]); r2.font.size = Pt(9)
    d.add_paragraph().paragraph_format.space_after = Pt(6)


def build_worksheet_docx(a, ws, out_path):
    d = new_doc()
    header_block(d, a, ws)

    line(d, "Record your team's output below.", bold=True, size=10.5, color=BRAND, after=8)

    n = len(ws["fields"])
    # keep every worksheet to 2 printed pages: fewer rules as the field count grows
    rules = 3 if n <= 8 else (2 if n <= 10 else 1)
    for idx, (num, label) in enumerate(ws["fields"]):
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.space_before = Pt(6)
        r = p.add_run(f"{num}.  "); r.bold = True; r.font.size = Pt(10.5)
        r.font.color.rgb = BRAND
        r2 = p.add_run(label); r2.bold = True; r2.font.size = Pt(10.5)
        answer_box(d, rules=rules)

    # self-check
    d.add_paragraph().paragraph_format.space_after = Pt(2)
    line(d, "SELF-CHECK BEFORE YOU FINISH", bold=True, size=10, color=TEAL, after=5)
    for ck in ws["checks"]:
        checkbox_line(d, ck)

    line(d, "", after=4)
    line(d, f"© 2026 {C.ORG}  ·  {C.UEN}  ·  Full step-by-step instructions: "
            f"Learner Guide, Activity {a['num']}",
         size=8, color=GREY, after=0, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_numbers(d)
    d.save(out_path)
    return out_path


def build_combined_docx(pairs, out_path):
    """One document holding every worksheet, for printing the whole set."""
    d = new_doc()
    # cover
    if os.path.exists(LOGO):
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(90)
        p.add_run().add_picture(LOGO, width=Inches(2.1))
    line(d, C.ORG, bold=True, size=13, after=1, align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, C.UEN, size=10, color=GREY, after=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, "ACTIVITY WORKSHEETS", bold=True, size=26, color=BRAND, after=6,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, "For", size=11, color=GREY, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, C.TITLE, bold=True, size=19, after=3, align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, f"TGS Ref No: {C.COURSE_CODE}", size=11, after=18,
         align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, f"{len(pairs)} activities  ·  one running case study: HarbourFront Logistics",
         size=10.5, color=GREY, after=6, align=WD_ALIGN_PARAGRAPH.CENTER)
    line(d, f"Version {C.VERSION}  ·  {C.VERSION_DATE}", bold=True, size=11, color=BRAND,
         after=0, align=WD_ALIGN_PARAGRAPH.CENTER)
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    for i, (a, ws) in enumerate(pairs):
        header_block(d, a, ws)
        line(d, "Record your team's output below.", bold=True, size=10.5, color=BRAND, after=8)
        rules = 3 if len(ws["fields"]) <= 8 else (2 if len(ws["fields"]) <= 10 else 1)
        for num, label in ws["fields"]:
            p = d.add_paragraph(); p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(6)
            r = p.add_run(f"{num}.  "); r.bold = True; r.font.size = Pt(10.5)
            r.font.color.rgb = BRAND
            r2 = p.add_run(label); r2.bold = True; r2.font.size = Pt(10.5)
            answer_box(d, rules=rules)
        d.add_paragraph().paragraph_format.space_after = Pt(2)
        line(d, "SELF-CHECK BEFORE YOU FINISH", bold=True, size=10, color=TEAL, after=5)
        for ck in ws["checks"]:
            checkbox_line(d, ck)
        if i < len(pairs) - 1:
            d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    add_page_numbers(d)
    d.save(out_path)
    return out_path


def to_pdf(docx_path, outdir):
    subprocess.run([SOFFICE, "--headless", "--convert-to", "pdf", "--outdir", outdir,
                    docx_path], capture_output=True, text=True, timeout=600)
    pdf = os.path.splitext(docx_path)[0] + ".pdf"
    return pdf if os.path.exists(pdf) else None


SLUGS = {
    1: "empathise-customer-design-thinking", 2: "diagnose-waterfall-fishbone",
    3: "product-backlog-sprint-planning", 4: "agile-raci-accountability",
    5: "execute-sprint-scrum-board", 6: "retrospective-5whys",
    7: "defect-pareto-analysis", 8: "velocity-forecast-metrics",
}

if __name__ == "__main__":
    if not os.path.exists(SOFFICE):
        raise SystemExit(f"LibreOffice not found at {SOFFICE}")
    pairs, made = [], []
    for a in ACTIVITIES:
        folder = os.path.join(ACTDIR, f"activity-{a['num']:02d}-{SLUGS[a['num']]}")
        md_path = os.path.join(folder, "WORKSHEET.md")
        if not os.path.exists(md_path):
            print(f"  SKIP activity {a['num']}: no WORKSHEET.md"); continue
        ws = parse_worksheet(open(md_path).read())
        pairs.append((a, ws))
        docx = build_worksheet_docx(a, ws, os.path.join(folder, "WORKSHEET.docx"))
        pdf = to_pdf(docx, folder)
        made.append((a["num"], len(ws["fields"]), len(ws["checks"]), pdf))
        print(f"  activity {a['num']}: {len(ws['fields'])} fields, "
              f"{len(ws['checks'])} checks -> WORKSHEET.pdf")

    combined = build_combined_docx(
        pairs, os.path.join(ACTDIR, f"Activity Worksheets - {C.SHORT_TITLE} - {C.VERSION}.docx"))
    cpdf = to_pdf(combined, ACTDIR)
    print(f"\nCombined: {os.path.basename(cpdf) if cpdf else 'FAILED'}")
    print(f"Worksheet PDFs written: {sum(1 for m in made if m[3])}/{len(ACTIVITIES)}")
