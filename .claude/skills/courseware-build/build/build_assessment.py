#!/usr/bin/env python3
"""Build the WSQ assessment set for Agile Project Management for Business.

MIRRORS THE ORIGINAL PAPERS EXACTLY (pulled from the TMS course record):
  * Written Assessment (SAQ) — 7 open-ended questions, one per knowledge code K1..K7,
    60 minutes, open book.
  * Case Study (CS)          — 1 scenario + 3 questions covering A2/A5, A1/A3, A4/A6/A7,
    60 minutes, open book.

Only the CONTENT is rewritten — from THIS course's slides, Learner Guide and activities.
The instrument, question count, K/A mapping and timings are unchanged.

Page layout enforced by the house standard:
  page 1  cover page (names the instrument; NO version-control record)
  page 2  A: Trainee Information + B: Instructions to Candidate + Grading block, nothing else
  page 3+ the questions (WA) or the scenario then the questions (CS)

Outputs four DOCX into <repo>/assessment/. No PDFs — assessments are DOCX only.
Answer keys are TRAINER-ONLY and must never be uploaded to the LMS or GitHub.
"""
import os, sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH as AL, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import course_data as C
from prodoc import (add_cover_page, add_page_numbers, style_headings, _shade_cell,
                    BRAND, DARK, GREY)

LMS_URL = "https://lms-tms.tertiaryinfotech.com/"


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
OUT = os.path.join(REPO, "assessment")
SKILL_ASSETS = os.path.join(os.path.dirname(HERE), "assets")
LOGO = os.path.join(SKILL_ASSETS, "tertiary-infotech-logo.png")
VER = C.VERSION


# ============================================================ doc helpers
def new_doc(instrument):
    d = Document()
    st = d.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
    for s in d.sections:
        s.left_margin = s.right_margin = Inches(0.9)
        s.top_margin = s.bottom_margin = Inches(0.8)
    style_headings(d)
    add_cover_page(d, instrument, C.TITLE, VER, org_logo=LOGO, course_code=C.COURSE_CODE)
    return d


def line(d, text="", bold=False, size=11, color=DARK, after=6, align=None, italic=False):
    p = d.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    if align is not None: p.alignment = align
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.size = Pt(size); r.font.color.rgb = color
    return p


def add_hyperlink(p, url, text):
    r_id = p.part.relate_to(
        url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True)
    link = OxmlElement("w:hyperlink"); link.set(qn("r:id"), r_id)
    run = OxmlElement("w:r"); rPr = OxmlElement("w:rPr")
    for tag, val in (("w:sz", "22"), ("w:color", "0563C1")):
        el = OxmlElement(tag); el.set(qn("w:val"), val); rPr.append(el)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "single"); rPr.append(u)
    run.append(rPr)
    t = OxmlElement("w:t"); t.text = text; run.append(t)
    link.append(run); p._p.append(link)
    return link


def page2_block(d, instrument, duration="1 hour"):
    """Page 2 — Trainee Information + Instructions + Grading, and NOTHING else."""
    line(d, C.TITLE, bold=True, size=15, color=BRAND, after=2, align=AL.CENTER)
    line(d, instrument, bold=True, size=12.5, color=DARK, after=2, align=AL.CENTER)
    line(d, f"Course Code: {C.COURSE_CODE}", size=10, color=GREY, after=14, align=AL.CENTER)

    line(d, "A: Trainee Information", bold=True, size=12, after=6)
    for t in ("Trainee Name (as per NRIC): ______________________________________________",
              "Last 3 digits and alphabet of NRIC / FIN: _________________",
              "Date: __________________"):
        line(d, t, after=6)
    line(d, "", after=6)

    line(d, "B: Instructions to Candidate", bold=True, size=12, after=6)
    items = ["This is an individual exercise.",
             "This is an open-book assessment. You may refer to the course slides, the "
             "Learner Guide and any approved course materials only.",
             f"A total of {duration} is given to complete this assessment.",
             "Answer all questions in your own words. All questions are open-ended — there "
             "are no multiple-choice options.",
             "Write your answers in the space provided on this document.",
             None,
             "No discussion with other candidates is permitted. No photography or recording "
             "of the assessment script is allowed."]
    for i, s in enumerate(items, 1):
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(5)
        if s is None:
            r = p.add_run(f"{i}.  Submit your completed answers on this document, and upload "
                          f"the completed file to the LMS at ")
            r.font.size = Pt(11)
            add_hyperlink(p, LMS_URL, LMS_URL)
            p.add_run(".").font.size = Pt(11)
        else:
            p.add_run(f"{i}.  {s}").font.size = Pt(11)
    line(d, "", after=8)

    line(d, "C: Grading", bold=True, size=12, after=6)
    line(d, "This is a competency-based assessment. You will be assessed as Competent (C) or "
            "Not Yet Competent (NYC). To be assessed Competent you must satisfactorily address "
            "every question. A candidate assessed NYC may be re-assessed.", size=10.5, after=10)
    line(d, "For Official Use Only", bold=True, size=11.5, after=6)
    line(d, "Grade: _______  (C / NYC)", after=6)
    line(d, "Assessor Name: __________________________   Assessor NRIC: ________________", after=6)
    line(d, "Date: ___________________________________   Signature: ____________________", after=6)
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def answer_space(d, n=6):
    line(d, "Answer:", bold=True, size=10.5, color=GREY, after=4)
    for _ in range(n):
        line(d, "_________________________________________________________________________"
                "_______________", size=11, color=RGBColor(0xAA, 0xAA, 0xAA), after=7)
    line(d, "", after=6)


def official_use(d):
    line(d, "_______________________________________________________________________________"
            "_______", color=GREY, after=8)
    line(d, "For Official Use Only", bold=True, size=11.5, after=6)
    line(d, "Grade: _______  (C / NYC)", after=6)
    line(d, "Assessor Name: __________________________   Assessor NRIC: ________________", after=6)
    line(d, "Date: ___________________________________   Signature: ____________________", after=6)


def keytable(d, rows):
    """Marking table rendered as a REAL table (never wrapped ASCII columns).

    Starts on a fresh page so the header row can never be orphaned at the foot of
    the previous page, away from its body rows.
    """
    d.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    line(d, "Marking Summary", bold=True, size=12, color=BRAND, after=8)
    t = d.add_table(rows=0, cols=3); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.add_row().cells
    for i, h in enumerate(("Q", "Code(s)", "What a competent answer must contain")):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hdr[i], "1F6FEB")
    trPr = t.rows[0]._tr.get_or_add_trPr()
    hdr_el = OxmlElement("w:tblHeader"); hdr_el.set(qn("w:val"), "true"); trPr.append(hdr_el)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            for j, ln in enumerate(str(v).split("\n")):
                p = cells[i].paragraphs[0] if j == 0 else cells[i].add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                rr = p.add_run(ln); rr.font.size = Pt(9)
                if i < 2: rr.bold = True
    for r_ in t.rows:
        r_.cells[0].width = Inches(0.4)
        r_.cells[1].width = Inches(0.9)
        r_.cells[2].width = Inches(5.4)
        # keep a marking row from splitting across a page break
        trPr = r_._tr.get_or_add_trPr()
        el = OxmlElement("w:cantSplit"); trPr.append(el)
    d.add_paragraph()


# ============================================================ WRITTEN — 7 Qs, K1..K7
WRITTEN = [
    ("K1",
     "A mid-sized Singapore logistics provider is deciding whether to move from a traditional "
     "waterfall approach to Agile, after an 11-month portal project was delivered 4 months late "
     "with most of its features unused. You are on the transition team and must assess how the "
     "change would affect current and future operations.",
     "What factors should you consider in assessing the value proposition of Agile compared to "
     "waterfall in this scenario, and what methods would you use to analyse the current and "
     "future business operating landscape?",
     ["Name methods for analysing the operating landscape: PESTLE, SWOT, Porter's Five Forces, "
      "scenario planning — and note that Agile re-runs these periodically rather than once.",
      "Identify the VUCA conditions (volatility, uncertainty, complexity, ambiguity) that make a "
      "frozen 11-month plan unreliable.",
      "Compare the value profile: waterfall delivers all value at the end; Agile releases value "
      "each sprint, so benefit is realised earlier and the return profile improves.",
      "Compare the cost of change: it compounds over time in waterfall, whereas short feedback "
      "loops flatten the curve because less is built on a wrong assumption.",
      "Explain the inverted iron triangle — waterfall fixes scope and flexes time/cost; Agile "
      "fixes time and cost and varies scope.",
      "Note the risk profile: waterfall concentrates risk in the final third (one late "
      "integration, compressed testing); Agile retires risk continuously.",
      "Acknowledge honestly what Agile does NOT fix — an unproven platform still needs a spike, "
      "and stage-gate funding governance must itself change or the adoption becomes theatre.",
      "Conclude with the decision criteria: stable, regulated scope favours waterfall; evolving "
      "requirements and high novelty favour Agile; many organisations run a deliberate hybrid."]),

    ("K2",
     "A software team is building a shipment-tracking portal for logistics customers. The "
     "previous version was built from a 96-page signed-off requirements document and 62% of its "
     "features go unused. The team must now ensure the product meets both current and future "
     "customer needs.",
     "What key considerations and methods should the team use when analysing customer needs and "
     "preferences throughout the project lifecycle?",
     ["State the core problem: a one-off requirements interview captures what people could "
      "imagine at the start, not what they need at delivery.",
      "Name the methods: personas and empathy maps (says / thinks / does / feels), journey maps, "
      "jobs-to-be-done, Kano analysis, and A/B experiments to test preference with behaviour.",
      "Explain reframing a requirement into a customer-value problem statement — a specific user, "
      "a specific need, and the consequence of not meeting it, with no solution embedded.",
      "Distinguish a feature description from a customer benefit: 'configurable export formats' "
      "is a feature; 'know if it clears customs today so I can re-sequence production' is value.",
      "Make analysis continuous: the Sprint Review turns customer contact into a rhythm every "
      "sprint, rather than a single event at UAT.",
      "Use the customer insight to ORDER the backlog by value, so the highest-value features are "
      "delivered first rather than whatever is easiest.",
      "Validate cheaply before building — wireframes and low-fidelity prototypes clarify what "
      "'done' looks like and test the approach before commitment."]),

    ("K3",
     "A multinational corporation is transitioning to Agile. You are responsible for aligning "
     "the organisation's policies, processes and standards with Agile ways of working. The "
     "existing governance model approves funding at stage gates and reports progress as "
     "percent-complete of tasks.",
     "How should organisations typically align their policies, processes and standards to support "
     "Agile methodologies?",
     ["Shift governance from stage-gate approval to incremental funding — the ability to extend, "
      "redirect or cancel at a granular level.",
      "Change the progress measure: working, usable output replaces percent-complete of tasks "
      "(Agile principle 7).",
      "Replace formal change-control-by-exception with a continuously re-ordered product backlog, "
      "so change is a normal input rather than a variance.",
      "Make the quality standard explicit and testable through a Definition of Done, so "
      "'complete' means the same thing to everyone and is auditable.",
      "Align procurement and contracting: fix time and cost, vary scope; consider graduated "
      "fixed-price or fixed-price work packages, and 'money for nothing, change for free'.",
      "Restructure roles and decision rights so one Product Owner is accountable for backlog "
      "order — duplicated accountability is the most common cause of Agile failure.",
      "Preserve compliance obligations: regulatory requirements are scheduled as real backlog "
      "items inside sprints, never deferred to the end.",
      "Adopt evolutionary rather than revolutionary change, so the organisation can absorb it — "
      "policies, processes and people move together or the adoption stalls."]),

    ("K4",
     "Your company is adopting Agile practices across all departments and you are managing the "
     "transition. Sponsors are worried about visible failure, middle managers fear losing "
     "control, and two delivery teams have openly resisted the change.",
     "What change management methodologies, tools and practices would be most effective in "
     "managing the risks of this organisational change?",
     ["Map the stakeholders and name each group's actual fear, then answer it directly rather "
      "than generically.",
      "Sponsors fear visible failure — answer: incremental funding allows stopping early and "
      "cheaply, which lowers exposure.",
      "Middle managers fear loss of control — answer: they gain real progress data (working "
      "output, cycle time) instead of estimated percent-complete.",
      "Teams fear exposure and surveillance — answer: the board makes the WORK visible, not "
      "individual performance.",
      "Users fear losing promised features — answer: they receive the highest-value features "
      "sooner and get a say every sprint.",
      "Finance and the PMO fear unauditable scope — answer: fixed cadence, fixed cost and "
      "auditable increments.",
      "Use practices that de-risk adoption: start with a willing pilot team on a real project, "
      "make wins visible with evidence, use evolutionary change, and coach for 2–3 sprints then "
      "withdraw.",
      "Manage delivery risk with Agile risk tools: a risk-adjusted backlog, a risk burndown "
      "chart, expected monetary value (probability × impact), and timeboxed spikes.",
      "State the honest warning: never sell Agile as a way to get more output for less money. It "
      "buys responsiveness, earlier value and lower risk of building the wrong thing."]),

    ("K5",
     "In preparation for launching a new product, your company is forming an Agile team of seven "
     "people drawn from different departments. Some members have never worked together, and each "
     "skill is currently held by only one person.",
     "Describe the different models of team composition and team development that could be "
     "applied in this context, and analyse their potential impact on the team's performance.",
     ["Composition: cross-functional (every skill needed to reach Done sits inside the team), "
      "self-organising, fewer than about 12 people, and stable in membership.",
      "Explain why stability matters — velocity and any forecast built on it are meaningless if "
      "the team changes every sprint.",
      "Identify the single-specialist risk: one person per skill guarantees queues and a "
      "bottleneck; generalising specialists (deep in one skill, capable across several) remove it.",
      "Tuckman's five stages — forming, storming, norming, performing, adjourning — with the "
      "storming dip being normal and necessary rather than a failure.",
      "Adaptive leadership: match style to stage — directing, then coaching, then supporting, "
      "then delegating. Delegating to a forming team fails; directing a performing team insults it.",
      "Shu-Ha-Ri and the Dreyfus model of skill acquisition: follow the practice, adapt the "
      "practice, then transcend it. New teams should follow the framework before tailoring it.",
      "Servant leadership as the operating style — shield the team, remove impediments, "
      "re-communicate the vision, and tap intrinsic motivation (autonomy, mastery, purpose).",
      "Impact on performance: co-location or deliberate connection enables osmotic communication; "
      "a safe environment for constructive disagreement produces better decisions and real buy-in."]),

    ("K6",
     "Your management team wants to ensure the organisation's new Agile processes genuinely "
     "align with the core values and principles of the Agile Manifesto, rather than simply "
     "renaming existing meetings.",
     "What are the core values and guiding principles of the Agile Manifesto, and how can they "
     "be effectively integrated into the company's culture?",
     ["State the four values as preferences, not prohibitions: individuals and interactions over "
      "processes and tools; working software over comprehensive documentation; customer "
      "collaboration over contract negotiation; responding to change over following a plan.",
      "Read the values correctly — 'we value the left over the right' means both matter, and the "
      "left matters more when they conflict. It does not mean the right is worthless.",
      "Summarise the twelve principles: early and continuous delivery; welcome late change; "
      "deliver frequently; business and delivery people work together daily; motivated "
      "individuals; face-to-face conversation.",
      "Continue: working output is the primary measure of progress; sustainable pace; technical "
      "excellence; simplicity (maximise work not done); self-organising teams; regular reflection "
      "and adjustment.",
      "Integration into culture: make the values visible in decisions — who is allowed to decide, "
      "what gets reported, and what happens when someone raises a problem.",
      "Give behavioural tests rather than slogans: is documentation barely sufficient and just in "
      "time? Is change managed rather than suppressed? Is progress shown as working output?",
      "Use principle 12 as the mechanism — the retrospective is where the values are actually "
      "enforced, with each improvement given an owner and capacity in the next sprint.",
      "Warn against cargo-cult adoption: holding stand-ups without changing how decisions are "
      "made and funded is ceremony, not agility."]),

    ("K7",
     "An IT consultancy firm is exploring several Agile methodologies to improve project "
     "outcomes for clients whose work ranges from new product development to a high-variability "
     "support queue. You are on the task force analysing the options.",
     "What Agile methodologies are available, and how do their practices differ in terms of "
     "process and team management?",
     ["Scrum — the most widely adopted (roughly 63% of Agile teams): 3 accountabilities (Product "
      "Owner, Scrum Master, Developers), 3 artefacts (Product Backlog, Sprint Backlog, "
      "Increment) and 5 events, built on empiricism. Fixed sprints of 1–4 weeks with a Sprint Goal.",
      "Kanban — visualise the workflow, limit WIP, manage flow, make policies explicit, improve "
      "collaboratively. Continuous flow with no prescribed roles; suits support queues and "
      "high-variability work.",
      "Explain Little's Law as Kanban's engine: Cycle Time = WIP ÷ Throughput, so limiting WIP "
      "reduces cycle time without working harder.",
      "Scrumban — a Scrum cadence combined with Kanban WIP limits, for teams outgrowing strict "
      "Scrum.",
      "Lean — seven principles from Toyota (eliminate waste, amplify learning, decide as late as "
      "possible, deliver fast, empower the team, build quality in, optimise the whole) targeting "
      "the eight wastes; waiting usually dominates in knowledge work.",
      "Extreme Programming (XP) — five values plus engineering practices: TDD, pair programming, "
      "continuous integration, refactoring, collective code ownership, small releases, "
      "sustainable pace.",
      "DSDM / AgilePM — the APM-endorsed framework with feasibility, foundations, evolutionary "
      "development and deployment; strong on governance, common in regulated delivery.",
      "Others: Feature-Driven Development, Crystal (sized by team and criticality), and "
      "SAFe / LeSS for scaling across many teams.",
      "State the selection rule: choose by the shape of the work, not by fashion — a support team "
      "forced into fixed sprints will fight the framework every week. Key metrics differ too: "
      "velocity and sprint burndown for Scrum; cycle time, throughput and the CFD for Kanban."]),
]

# ============================================================ CASE STUDY — 1 scenario + 3 Qs
CS_SCENARIO_TITLE = "Scenario: HarbourFront Logistics — the CustomerConnect restart"

CS_SCENARIO = [
    "HarbourFront Logistics Pte Ltd is a Singapore third-party logistics (3PL) provider with 140 "
    "staff. It handles inbound sea and air freight for mid-sized manufacturers, manages customs "
    "clearance, and operates a bonded warehouse in Jurong. Its customers depend on predictable "
    "inbound flow to keep their production lines running.",

    "In January 2024 HarbourFront delivered 'CustomerConnect', a self-service portal intended to "
    "let customers track their own shipments. It ran as a waterfall project: a 96-page "
    "requirements document was signed off in month 1 and frozen, the customer was consulted at "
    "requirements and again at UAT in month 13, the components were first integrated in month 9, "
    "and testing was compressed into the final six weeks. It was delivered 4 months late. Of the "
    "24 features shipped, only 3 are used weekly. Inbound call volume is unchanged, and the top "
    "complaint to the service desk is still \"where is my shipment right now?\"",

    "The management committee has approved an Agile restart: 6 sprints of 2 weeks, one team of "
    "7 people, and a demonstration to the top 5 customers in 12 weeks. You have been appointed "
    "Agile project lead.",

    "Four sprints in, the position is as follows. The team has completed sprints with velocities "
    "of 14, 18, 17 and 21 story points, and 186 points remain in the backlog. Average cycle time "
    "has risen from 3.1 days in Sprint 1 to 7.4 days in Sprint 4. On the cumulative flow diagram, "
    "the Testing band widens steadily over time. 120 defects have been logged, of which 44 are "
    "stale or missing feed data, 31 unclear acceptance criteria, 17 environment drift, 11 carrier "
    "API changes, and the remaining 17 spread across four minor categories.",

    "Three organisational problems persist. Information about a critical feature sits only with "
    "the design specialist, blocking the two developers. The legacy programme director still "
    "approves every story change and asks for a weekly percent-complete report. The operations "
    "manager gives the team priorities directly, so developers receive conflicting instructions. "
    "The team has one person per skill.",
]

CS_QUESTIONS = [
    ("A2, A5",
     "During a sprint planning session you notice that information about a critical feature "
     "requirement is held only by the design specialist, creating a bottleneck that is delaying "
     "the two developers. Analyse the current communication and work-organisation processes at "
     "HarbourFront, and recommend how to foster a cross-functional Agile mindset that keeps "
     "information flowing across the team. In your response, explain how an Agile mindset would "
     "help the team organise work in alignment with priorities, and how you would encourage the "
     "team to experiment with new collaborative methods.",
     ["Diagnose the structural cause: one specialist per skill guarantees a queue behind the only "
      "person who can do the work — this is a design flaw in the team, not a personal failing.",
      "Recommend generalising specialists and deliberate cross-skilling (pairing, shadowing, "
      "collective ownership) so more than one person can progress the work.",
      "Make the information visible rather than personal: the board, an explicit Definition of "
      "Done, acceptance criteria written as Given/When/Then, and information radiators.",
      "Use the Scrum events as the communication mechanism — the Daily Scrum surfaces the "
      "impediment within 24 hours, and it is raised the same day rather than on day 9.",
      "Organising work in alignment with priorities (A2): one ordered product backlog, a single "
      "Sprint Goal, and a WIP limit so the team finishes before it starts more.",
      "Reference the RACI work from Activity 4 — make accountability explicit so exactly one role "
      "is Accountable for backlog order, which removes the conflicting-priority problem.",
      "Experimenting with new methods (A5): create a safe environment to try a change for one "
      "sprint, measure it, and keep or discard it at the retrospective — for example pairing, a "
      "lower WIP limit, or a shared feature-definition workshop.",
      "Close the loop with PDCA: each experiment becomes a retrospective action with a named "
      "owner and story points in the next sprint backlog, not a good intention."]),

    ("A1, A3",
     "The product backlog for the CustomerConnect restart is extensive and must be ordered for "
     "the next sprint. Using the defect data and the customer insight available to you, analyse "
     "the backlog and determine how you would prioritise it to reduce waste and improve "
     "operational efficiency. Explain the prioritisation techniques you would use and how you "
     "would share the resulting decisions across the team and with stakeholders.",
     ["Prioritise by customer value first, using the reframed customer need — live customs and "
      "delivery status is what stops Priya's production line, and it is still unmet.",
      "Name and apply prioritisation techniques: MoSCoW (and be strict — if more than 60% is "
      "'Must', nothing has been prioritised), dot voting, the 100-point method, Kano analysis, "
      "and weighted shortest job first.",
      "Use the Pareto analysis to reduce waste (A3): the top 3 defect causes account for 76.7% of "
      "all 120 defects, and the top 4 for 85.8%. Fixing 2 of 8 causes addresses 62.5%.",
      "Connect the root cause from the 5 Whys work: a Definition of Done copied from the previous "
      "project with no data-quality check sits underneath the two largest defect categories, so "
      "one fix attacks both.",
      "Schedule the root-cause fix and the technical-debt work as real backlog items with story "
      "points — improvement with no allocated capacity does not happen.",
      "Consider technical dependencies and team capacity: enabling stories (carrier feed, "
      "authentication, alerting) must be ordered so each sprint can still produce a demonstrable "
      "increment; split by workflow step, never by technical layer.",
      "Apply INVEST and split anything estimated at 13 or 21 points — a number that large means "
      "the team does not yet understand the item.",
      "Sharing information across teams (A1): the ordered backlog is visible to everyone, the "
      "Sprint Review demonstrates the increment to stakeholders, and the trade-off is stated "
      "explicitly — fixing all 8 causes would cost roughly 3 sprints to eliminate a final 14.2%.",
      "Recommend one accountable decision-maker for order (the Product Owner), with the "
      "operations manager as Consulted rather than a second Accountable."]),

    ("A4, A6, A7",
     "The management committee meets on Friday and wants to know when the remaining "
     "CustomerConnect scope will be delivered. Separately, the operations manager has asked why "
     "the support queue keeps growing even though the team appears to be going faster. Analyse "
     "how the team currently measures progress against its sprint goals and business objectives, "
     "and recommend a set of Agile metrics and practices that would track progress more "
     "accurately, support continuous improvement, and build individual and team accountability.",
     ["Build the forecast from empirical data (A4): average velocity is 17.5, so 186 ÷ 17.5 ≈ 11 "
      "sprints; the slowest sprint (14) gives 14 sprints and the fastest (21) gives 9.",
      "Present it as a RANGE with stated assumptions — roughly 18 to 28 weeks, most likely about "
      "22 weeks, assuming a stable team, no backlog growth, the same estimation scale and no lost "
      "sprints. Refuse to give a single date.",
      "Recommend the five metrics and what each is for: sprint burndown (progress within the "
      "sprint), release/epic burndown (progress across the release), velocity (forecasting), the "
      "control chart (cycle and lead time) and the cumulative flow diagram (locating the "
      "bottleneck).",
      "Diagnose the growing support queue: the widening Testing band on the CFD identifies the "
      "bottleneck, and the control chart confirms it — cycle time rose from 3.1 to 7.4 days.",
      "Explain it with Little's Law: Cycle Time = WIP ÷ Throughput. Velocity rose because more "
      "work was STARTED, while throughput past the Testing constraint did not, so WIP and cycle "
      "time both rose. The team looks faster and delivers to the customer more slowly.",
      "Recommend the counter-intuitive action: lower the WIP limit on Testing and stop starting "
      "new stories until Testing clears. Increasing velocity would make it worse.",
      "State the metric anti-pattern explicitly: velocity is a forecasting input, never a target. "
      "Rewarding it produces estimate inflation and destroys the forecast; never compare velocity "
      "between teams.",
      "Continuous improvement (A6): use the Sprint Review to inspect the product with "
      "stakeholders and the Retrospective to inspect the process, running the five stages and "
      "using 5 Whys, fishbone and Pareto to reach a cause the team can change.",
      "Insist that 'Done' means Done — work failing the Definition of Done returns to the backlog "
      "at full estimate and is never counted as velocity, otherwise every forecast is fiction.",
      "Accountability (A7): the team commits collectively to the Sprint Goal, each member pulls "
      "their own work and updates the board honestly, impediments are raised the same day, and "
      "each improvement action carries a named owner and story points."]),
]


# ============================================================ builders
def build_written_paper():
    d = new_doc("Written Assessment (SAQ)")
    page2_block(d, "Written Assessment (SAQ)", "1 hour")
    line(d, "D: Short-Answer Questions", bold=True, size=12, after=4)
    line(d, f"Answer all {len(WRITTEN)} questions. All questions are open-ended.",
         size=10.5, italic=True, color=GREY, after=12)
    for i, (code, ctx, q, _ans) in enumerate(WRITTEN, 1):
        line(d, f"Question {i}.", bold=True, size=11.5, after=4)
        line(d, ctx, size=10.5, italic=True, color=GREY, after=5)
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(q); r.font.size = Pt(11)
        r2 = p.add_run(f"  ({code})"); r2.font.size = Pt(11); r2.bold = True
        r2.font.color.rgb = BRAND
        answer_space(d, 7)
    official_use(d)
    add_page_numbers(d)
    out = os.path.join(OUT, f"WA (SAQ) - {C.SHORT_TITLE} - {VER}.docx")
    d.save(out); return out


def build_written_answers():
    d = new_doc("Written Assessment (SAQ) — Answer Key")
    line(d, C.TITLE, bold=True, size=15, color=BRAND, after=2, align=AL.CENTER)
    line(d, "Written Assessment (SAQ) — Model Answers and Marking Guide", bold=True,
         size=12.5, after=2, align=AL.CENTER)
    line(d, f"Course Code: {C.COURSE_CODE}   ·   TRAINER COPY — NOT FOR LEARNERS",
         size=10, color=GREY, after=14, align=AL.CENTER)
    line(d, f"This paper carries {len(WRITTEN)} questions, one for each knowledge code "
            f"K1–K7, mirroring the original paper. Suggested answers are not exhaustive; "
            f"credit any technically correct response that addresses the code.",
         size=10.5, italic=True, after=12)
    for i, (code, ctx, q, ans) in enumerate(WRITTEN, 1):
        line(d, f"Question {i}  ({code})", bold=True, size=11.5, color=BRAND, after=4)
        line(d, q, size=10.5, italic=True, color=GREY, after=5)
        line(d, "Suggestive answers (not exhaustive):", bold=True, size=10.5, after=4)
        for a in ans:
            p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
            p.add_run(a).font.size = Pt(10.5)
        line(d, "", after=8)
    keytable(d, [[str(i), code,
                  f"Addresses {code}. Competent requires at least 4 of the listed points, "
                  f"in the candidate's own words."]
                 for i, (code, _c, _q, _a) in enumerate(WRITTEN, 1)])
    line(d, "A candidate is assessed Competent on this instrument only when every question "
            "is satisfactorily addressed.", size=10.5, bold=True, after=6)
    add_page_numbers(d)
    out = os.path.join(OUT, f"Answer to WA (SAQ) - {C.SHORT_TITLE} - {VER}.docx")
    d.save(out); return out


def build_case_paper():
    d = new_doc("Case Study (CS)")
    page2_block(d, "Case Study (CS) Assessment", "1 hour")
    line(d, "D: Case Study", bold=True, size=12, after=6)
    line(d, CS_SCENARIO_TITLE, bold=True, size=11.5, color=BRAND, after=6)
    for pgh in CS_SCENARIO:
        line(d, pgh, size=10.5, after=6)
    line(d, "", after=6)
    line(d, f"Answer all {len(CS_QUESTIONS)} questions below, with reference to the scenario "
            f"above. All questions are open-ended.", size=10.5, italic=True, color=GREY, after=10)
    for i, (code, q, _a) in enumerate(CS_QUESTIONS, 1):
        line(d, f"Question {i}.", bold=True, size=11.5, after=4)
        p = d.add_paragraph(); p.paragraph_format.space_after = Pt(6)
        r = p.add_run(q); r.font.size = Pt(11)
        r2 = p.add_run(f"  ({code})"); r2.font.size = Pt(11); r2.bold = True
        r2.font.color.rgb = BRAND
        answer_space(d, 10)
    official_use(d)
    add_page_numbers(d)
    out = os.path.join(OUT, f"CS Assessment - {C.SHORT_TITLE} - {VER}.docx")
    d.save(out); return out


def build_case_answers():
    d = new_doc("Case Study (CS) — Answer Key")
    line(d, C.TITLE, bold=True, size=15, color=BRAND, after=2, align=AL.CENTER)
    line(d, "Case Study (CS) — Model Answers and Marking Guide", bold=True, size=12.5,
         after=2, align=AL.CENTER)
    line(d, f"Course Code: {C.COURSE_CODE}   ·   TRAINER COPY — NOT FOR LEARNERS",
         size=10, color=GREY, after=14, align=AL.CENTER)
    line(d, f"This paper carries one scenario and {len(CS_QUESTIONS)} questions covering the "
            f"ability codes A1–A7, mirroring the original paper. The scenario is the same "
            f"HarbourFront Logistics case used in the in-class activities, so every model answer "
            f"traces to work the learners actually did.", size=10.5, italic=True, after=10)
    line(d, CS_SCENARIO_TITLE, bold=True, size=11.5, color=BRAND, after=6)
    for pgh in CS_SCENARIO:
        line(d, pgh, size=10, color=GREY, after=5)
    line(d, "", after=8)
    ACT_TRACE = {
        1: "Activities 1, 4 and 5 (Design Thinking reframing, RACI accountability, sprint execution)",
        2: "Activities 3, 6 and 7 (backlog and sprint planning, 5 Whys root cause, Pareto analysis)",
        3: "Activities 5, 6 and 8 (sprint tracking, retrospective, velocity forecast and metrics)",
    }
    for i, (code, q, ans) in enumerate(CS_QUESTIONS, 1):
        line(d, f"Question {i}  ({code})", bold=True, size=11.5, color=BRAND, after=4)
        line(d, q, size=10, italic=True, color=GREY, after=5)
        line(d, f"Taught and practised in: {ACT_TRACE[i]}", size=10, bold=True, after=5)
        line(d, "Suggestive answers (not exhaustive):", bold=True, size=10.5, after=4)
        for a in ans:
            p = d.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
            p.add_run(a).font.size = Pt(10.5)
        line(d, "", after=8)
    keytable(d, [[str(i), code,
                  f"Addresses {code}. Competent requires at least 5 of the listed points, "
                  f"applied to the HarbourFront scenario rather than stated generically.\n"
                  f"Traces to: {ACT_TRACE[i]}"]
                 for i, (code, _q, _a) in enumerate(CS_QUESTIONS, 1)])
    line(d, "A candidate is assessed Competent on this instrument only when all three questions "
            "are satisfactorily addressed with reference to the scenario.", size=10.5, bold=True,
         after=6)
    add_page_numbers(d)
    out = os.path.join(OUT, f"Answer to CS Assessment - {C.SHORT_TITLE} - {VER}.docx")
    d.save(out); return out


if __name__ == "__main__":
    os.makedirs(OUT, exist_ok=True)
    files = [build_written_paper(), build_written_answers(),
             build_case_paper(), build_case_answers()]
    for f in files:
        print("Saved:", f)

    # ---- self-verify against the ORIGINAL papers pulled from the TMS ----
    wk = [w[0] for w in WRITTEN]
    ck = []
    for code, _q, _a in CS_QUESTIONS:
        ck += [c.strip() for c in code.split(",")]
    print()
    print(f"WA questions: {len(WRITTEN)}  (original: 7)")
    print(f"WA K codes:   {wk}")
    print(f"CS questions: {len(CS_QUESTIONS)}  (original: 3)")
    print(f"CS A codes:   {ck}")
    assert len(WRITTEN) == 7, "WA question count must mirror the original (7)"
    assert len(CS_QUESTIONS) == 3, "CS question count must mirror the original (3)"
    assert wk == [f"K{i}" for i in range(1, 8)], "WA must cover K1..K7, one question each"
    assert sorted(set(ck)) == [f"A{i}" for i in range(1, 8)], \
        f"CS must cover A1..A7, got {sorted(set(ck))}"
    print("\nAll K1-K7 covered by the WA. All A1-A7 covered by the CS.")
    print("Question counts and instrument types mirror the originals.")
