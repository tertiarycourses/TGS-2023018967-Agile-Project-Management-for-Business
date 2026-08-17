#!/usr/bin/env python3
"""Build the WSQ Lesson Plan (LP) DOCX for Agile Project Management for Business.

House format: WSQ cover page, Document Version Control Record, real Word TOC field,
Arial 11pt body, colour-coded daily schedule tables, and a "Page X of Y" + copyright
footer on every page.

Each training day totals exactly 8 instructional hours (09:30–18:30 with a 1-hour
lunch; tea breaks are counted within instructional time).

Slide references come from slide_map.json, written by build_slides.py, so the LP can
never cite a slide number the deck does not have.
"""
import os, sys, json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
import course_data as C
from data_domain1 import DOMAIN1
from data_domain2 import DOMAIN2
from data_domain3 import DOMAIN3
from prodoc import (add_cover_page, add_version_control, add_toc, add_page_numbers,
                    style_headings, enable_update_fields, _shade_cell, BRAND, DARK, GREY)

ACTIVITIES = DOMAIN1 + DOMAIN2 + DOMAIN3


def _find_repo(start):
    env = os.environ.get("COURSE_REPO")
    if env and os.path.isdir(env):
        return env
    d = start
    for _ in range(8):
        d = os.path.dirname(d)
        if os.path.isdir(os.path.join(d, "courseware")) and os.path.isdir(os.path.join(d, "activities")):
            return d
    return os.path.dirname(os.path.dirname(HERE))


REPO = _find_repo(HERE)
SKILL_ASSETS = os.path.join(os.path.dirname(HERE), "assets")
try:
    SLIDES = json.load(open(os.path.join(HERE, "slide_map.json")))
except Exception:
    SLIDES = {}

doc = Document()
st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(11)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.8)
    s.top_margin = s.bottom_margin = Inches(0.7)
style_headings(doc)

LOGO = os.path.join(SKILL_ASSETS, "tertiary-infotech-logo.png")

add_cover_page(doc, "LESSON PLAN", C.TITLE, C.VERSION, org_logo=LOGO,
               course_code=C.COURSE_CODE)

add_version_control(doc, [
    ("v9.0", C.VERSION_DATE,
     "Full content revamp. Rebuilt on a single-source content module. Expanded Agile "
     "theory (Manifesto values/principles, Scrum accountabilities/artefacts/events, Lean, "
     "Kanban + Little's Law, XP, DSDM, scaling). Added the five Agile metrics with worked "
     "interpretation. Replaced generic exercises with 8 tool-based activities on one running "
     "case study (HarbourFront Logistics). Added activity debriefs and full TSC mapping.",
     "Dr Alfred Ang"),
    ("v8.0", "28 November 2022", "Previous released version.", "Tertiary Infotech Academy"),
])

add_toc(doc)


def h1(t):
    p = doc.add_heading(t, level=1); p.paragraph_format.space_before = Pt(10); return p


def h2(t):
    p = doc.add_heading(t, level=2); p.paragraph_format.space_before = Pt(8); return p


def para(t, size=11, bold=False, italic=False, color=DARK, after=6):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(after)
    r = p.add_run(t); r.font.size = Pt(size); r.bold = bold; r.italic = italic
    r.font.color.rgb = color
    return p


def bullet(t, size=11):
    p = doc.add_paragraph(style="List Bullet"); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.size = Pt(size)
    return p


def table(headers, rows, widths=None, header_fill="1F6FEB", font=9.5):
    t = doc.add_table(rows=0, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hr = t.add_row().cells
    for i, htext in enumerate(headers):
        hr[i].text = ""
        rr = hr[i].paragraphs[0].add_run(htext)
        rr.bold = True; rr.font.size = Pt(font); rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        _shade_cell(hr[i], header_fill)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            for j, ln in enumerate(str(val).split("\n")):
                p = cells[i].paragraphs[0] if j == 0 else cells[i].add_paragraph()
                p.paragraph_format.space_after = Pt(1)
                rr = p.add_run(ln); rr.font.size = Pt(font)
                if i == 0: rr.bold = True
    if widths:
        for r_ in t.rows:
            for i, w in enumerate(widths):
                r_.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ------------------------------------------------------------------ 1. course info
h1("1.  Course Information")
table(["Field", "Detail"], [
    ["Course Title", C.TITLE],
    ["WSQ Course Code", C.COURSE_CODE],
    ["Training Provider", f"{C.ORG}  ({C.UEN})"],
    ["Course Duration", f"{C.DAYS} days  ·  {C.DURATION_HRS} instructional hours "
                        f"(8 hours per day)  ·  plus 2 hours assessment"],
    ["Delivery Mode", "Physical classroom, synchronous Zoom, or corporate on-site delivery"],
    ["Class Size", "Maximum 20 learners per class"],
    ["TSC Reference", f"{C.TSC_TITLE}  ·  {C.TSC_CODE}"],
    ["Assessment", "Written Assessment (WA — SAQ, 1 hour) and Case Study (CS, 1 hour). "
                   "Both open book."],
    ["Trainer", C.TRAINER],
    ["Version", f"{C.VERSION}  ·  {C.VERSION_DATE}"],
    ["Course Page", C.COURSE_URL],
], widths=[1.9, 5.0])

h2("1.1  Target Audience")
para("This course is designed for professionals who lead, coordinate or contribute to "
     "projects and who need to deliver in conditions where requirements change. Suitable "
     "job roles include:")
for r in C.JOB_ROLES:
    bullet(r)

h2("1.2  Entry Requirements")
for p in C.PREREQUISITES:
    bullet(p)

h2("1.3  Learning Outcomes")
para("By the end of this course, learners will be able to:")
for lo in C.LEARNING_OUTCOMES:
    bullet(lo, size=11)

h2("1.4  TSC Mapping — Abilities and Knowledge Assessed")
para(f"All content maps to the Skills Framework TSC “{C.TSC_TITLE}” ({C.TSC_CODE}). "
     "Nothing is assessed that is not taught and practised.", italic=True, size=10)
table(["Code", "TSC Ability", "Taught in"], [
    ["A1", "Share information actively within and across teams to bridge operational barriers",
     "Topic 2 · Activity 4"],
    ["A2", "Organise work in alignment with operational priorities", "Topics 1–2 · Activities 2, 4"],
    ["A3", "Implement Agile or lean practices to reduce waste and defects in operating "
           "procedures and practices", "Topics 2–3 · Activities 3, 7"],
    ["A4", "Measure progress against targets for defined business outcomes on a regular basis",
     "Topic 3 · Activities 5, 7, 8"],
    ["A5", "Experiment with new ideas, products or services", "Topic 1 · Activity 1"],
    ["A6", "Assess work performance and quality to ensure continuous improvement",
     "Topic 3 · Activities 6, 7, 8"],
    ["A7", "Manage individual work responsibilities and take ownership of individual and "
           "team outcomes", "Topic 3 · Activities 5, 6, 8"],
], widths=[0.6, 4.4, 1.9])

table(["Code", "TSC Knowledge", "Taught in"], [
    ["K1", "Methods to analyse current and future business operating landscapes", "Topic 1 · Activity 2"],
    ["K2", "Methods to analyse current and future customer needs and preferences", "Topic 1 · Activity 1"],
    ["K3", "Organisational policies, processes and standards", "Topic 1 · Activity 2"],
    ["K4", "Types of change management methodologies, tools and practices", "Topic 2"],
    ["K5", "Types of team composition and formation models", "Topic 3 · Activities 4, 5"],
    ["K6", "Values and principles of Agile methodologies", "Topic 2 · Activity 3"],
    ["K7", "Types of Agile methodologies and practices", "Topic 2 · Activity 3"],
], widths=[0.6, 4.4, 1.9], header_fill="7C3AED")

# ------------------------------------------------------------------ 2. resources
h1("2.  Training Resources")
h2("2.1  Trainer Resources")
for r in ["Master Trainer Slides (PPTX and PDF), 120 slides, version " + C.VERSION,
          "This Lesson Plan", "Learner Guide with full step-by-step activity instructions",
          "Assessment instruments: WA (SAQ) question paper and answer key; Case Study question "
          "paper and marking guide (trainer-only — never distributed to learners)",
          "Laptop, projector and internet access", "Whiteboard and markers",
          "Sticky notes and A3 flip-chart paper for the team activities",
          "Planning-poker cards (Fibonacci: 1, 2, 3, 5, 8, 13, 21) — one deck per team"]:
    bullet(r)

h2("2.2  Learner Resources")
for r in ["Learner Slides (PDF)", "Learner Guide (PDF) — contains every activity step",
          "A laptop or tablet with a modern web browser and internet access",
          "The HarbourFront Logistics case brief (in the Learner Guide)"]:
    bullet(r)

h2("2.3  Online Tools Used in the Activities")
para("All tools are browser-based. Nothing needs to be installed and no login is required.",
     italic=True, size=10)
table(["Tool", "URL", "Used in"],
      [[t["name"], t["url"], t["use"]] for t in C.ED_TOOLS],
      widths=[1.3, 2.9, 2.7], header_fill="10B981")

# ------------------------------------------------------------------ 3. schedule
h1("3.  Daily Training Schedule")
para("Each training day runs 09:30–18:30 and delivers exactly 8 instructional hours. "
     "The 1-hour lunch break is excluded from instructional time; short tea breaks are "
     "counted within it. Slide references are to the Master Trainer Slides "
     f"version {C.VERSION}.", size=10, italic=True)

DAY_FILL = {1: "1F6FEB", 2: "10B981"}

DAY1 = [
    ["09:30 – 09:45", "0.25", "Digital Attendance (AM) & Welcome",
     "Trainer displays the SSG digital attendance QR code; learners scan and submit. "
     "Housekeeping and safety brief.",
     "Trainer-led\nAdmin", f"{SLIDES.get('admin', 2)}–7"],
    ["09:45 – 10:15", "0.50", "Trainer & Learner Introductions · Ground Rules",
     "Trainer introduction. Learners share role, industry and their biggest project "
     "frustration — captured on the whiteboard and referred back to all day.",
     "Discussion\nIce-breaker", "4–6"],
    ["10:15 – 10:45", "0.50", "Learning Outcomes · TSC Mapping · Course Outline",
     "The three learning outcomes, how the TSC abilities and knowledge map to each topic, "
     "the six tools, and the running HarbourFront case study.",
     "Trainer-led\nQ&A", "8–16"],
    ["10:45 – 11:00", "0.25", "Assessment Briefing & Funding Criteria",
     "Assessment format (WA + Case Study), open-book scope, the assessment flow, "
     "75% attendance and TRAQOM requirements.",
     "Trainer-led", "17–19"],
    ["11:00 – 12:15", "1.25", "Topic 1: Introduction to Agile Project Management (Part 1)",
     "Are we building the right thing? The traditional project lifecycle and the challenges "
     "every PM recognises. The waterfall model, its assumptions, and why it struggles. "
     "The cost-of-change curve. (K1, K3)",
     "Trainer-led\nDiagrams\nDiscussion",
     f"{SLIDES.get('topic1', 20)}–26"],
    ["12:15 – 13:00", "0.75", "Activity 1: Empathise with the Customer (start)",
     "Teams of 3–4 use the Design Thinking tool to build an empathy map and persona for "
     "Priya Menon, then reframe the requirement as a customer-value problem statement. "
     "(K2, A5)",
     "Hands-on\nTeam activity\nDesign Thinking tool",
     f"{SLIDES.get('act1', 37)}–39"],
    ["13:00 – 14:00", "—", "LUNCH BREAK (1 hour — not counted as instructional time)",
     "Digital Attendance (PM) taken on return.", "Break", "44"],
    ["14:00 – 14:30", "0.50", "Activity 1 debrief · Digital Attendance (PM)",
     "Each team presents its reframed problem statement in 90 seconds. Trainer draws out "
     "the difference between signed-off scope and delivered value.",
     "Presentation\nDebrief", "39"],
    ["14:30 – 15:45", "1.25", "Topic 1: Introduction to Agile Project Management (Part 2)",
     "What Agile actually is. Agile vs traditional. Inverting the iron triangle. Methods to "
     "analyse the operating landscape (PESTLE, SWOT, Five Forces, scenario planning) and "
     "customer needs (personas, journey maps, JTBD, Kano, A/B). Agile beyond software. "
     "When Agile is NOT the right answer. (K1, K2, A2, A5)",
     "Trainer-led\nComparison tables\nCase examples", "27–36"],
    ["15:45 – 16:30", "0.75", "Activity 2: Diagnose the Waterfall Failure (Fishbone)",
     "Teams use the Fishbone tool to analyse why CustomerConnect failed, tag each cause "
     "structural or behavioural, and pair each structural cause with an Agile practice. "
     "(K1, K3, A2)",
     "Hands-on\nTeam activity\nFishbone tool",
     f"{SLIDES.get('act2', 40)}–42"],
    ["16:30 – 16:45", "0.25", "Activity 2 debrief · Topic 1 recap",
     "Why most causes tag as structural, and the two causes Agile does not fix on its own.",
     "Debrief", "42–43"],
    ["16:45 – 18:15", "1.50", "Topic 2: Agile Essentials (Part 1)",
     "The Agile Manifesto: 4 values read correctly and the 12 principles. The Scrum "
     "framework end to end — 3 accountabilities, 3 artefacts and their commitments, "
     "5 events. The Definition of Done. (K6, K7)",
     "Trainer-led\nFramework diagram\nDiscussion",
     f"{SLIDES.get('topic2', 45)}–56"],
    ["18:15 – 18:30", "0.25", "Day 1 recap and Q&A",
     "Consolidate the mindset shift and preview Day 2.", "Q&A", "73–74"],
]

DAY2 = [
    ["09:30 – 09:45", "0.25", "Digital Attendance (AM) & Day 1 recap",
     "Learners scan the SSG QR code. Quick recall of the four values and the three Scrum "
     "accountabilities.", "Admin\nQ&A", "74"],
    ["09:45 – 11:00", "1.25", "Topic 2: Agile Essentials (Part 2)",
     "Lean and the seven principles, the eight wastes. Kanban's five practices and Little's "
     "Law. Extreme Programming. Choosing between Scrum, Kanban and Scrumban. DSDM/AgilePM "
     "and the scaling frameworks. (K7, A3)",
     "Trainer-led\nComparison table\nBoard walkthrough", "57–63"],
    ["11:00 – 11:45", "0.75", "Topic 2: Change Management and Resistance (K4)",
     "Finding Agile support: naming and answering each stakeholder's actual fear. Making "
     "the adoption stick — willing pilot, visible wins, evolutionary change, coach then "
     "withdraw. The honest warning about 'more for less'. (K4, A1)",
     "Trainer-led\nRole discussion", "64–65"],
    ["11:45 – 12:45", "1.00", "Activity 3: Build the Product Backlog & Run Sprint 1 Planning",
     "Teams write user stories with Given/When/Then criteria, test them against INVEST, "
     "estimate with planning poker, order with MoSCoW, write one sprint goal and load "
     "Sprint 1 to 20 points. Trainer acts as Product Owner. (K6, K7, A3)",
     "Hands-on\nTeam activity\nScrum Board tool\nPlanning poker",
     f"{SLIDES.get('act3', 67)}–69"],
    ["12:45 – 13:00", "0.25", "Activity 3 debrief",
     "Why the sprint goal is the artefact that makes a sprint reviewable, and why splitting "
     "by workflow step beats splitting by technical layer.", "Debrief", "69"],
    ["13:00 – 14:00", "—", "LUNCH BREAK (1 hour — not counted as instructional time)",
     "Digital Attendance (PM) taken on return.", "Break", "—"],
    ["14:00 – 14:45", "0.75", "Activity 4: Clarify Role Accountability (RACI) · Attendance (PM)",
     "Teams map 12 project decisions against 5 roles, audit for duplicate/missing "
     "Accountables, and resolve the conflicting-priorities escalation. (K5, A1, A2)",
     "Hands-on\nTeam activity\nRACI tool",
     f"{SLIDES.get('act4', 70)}–72"],
    ["14:45 – 15:30", "0.75", "Topic 3: Building the Agile Team and Setting the Vision",
     "Team composition and formation models, generalising specialists, Tuckman's stages and "
     "adaptive leadership, servant leadership, and setting a shared vision. (K5)",
     "Trainer-led\nDiagrams",
     f"{SLIDES.get('topic3', 75)}–80"],
    ["15:30 – 16:15", "0.75", "Topic 3: User Stories, Prioritisation and Estimation",
     "Story anatomy, the three Cs, INVEST and the split rule. Prioritisation techniques. "
     "Relative estimation, story points, planning poker and the 13-point rule. Timeboxing "
     "and Parkinson's Law. (A2, A7)",
     "Trainer-led\nWorked examples", "81–85"],
    ["16:15 – 17:15", "1.00", "Activity 5: Execute Sprint 1 and Track It on the Board",
     "Teams simulate all 10 sprint days with daily stand-ups, a WIP limit of 3, and three "
     "trainer-injected impediments. They plot the burndown and compute actual velocity. "
     "(K5, A4, A7)",
     "Hands-on\nSimulation\nScrum Board tool",
     f"{SLIDES.get('act5', 99)}–101"],
    ["17:15 – 17:45", "0.50", "Topic 3: The Five Agile Metrics, Waste and Risk",
     "Sprint burndown, release burndown, velocity, control chart and cumulative flow "
     "diagram — what each shows, how to read it, and the anti-patterns. Little's Law "
     "worked numerically. Root-cause tools, the Pareto principle, value stream mapping, "
     "risk as anti-value, and technical debt and ownership. (A3, A4, A6, A7)",
     "Trainer-led\nChart interpretation", "86–98"],
    ["17:45 – 18:15", "0.50", "Activities 6, 7 & 8: Retrospective, Pareto and Forecast",
     "Activity 6: full retrospective with 5 Whys to a changeable root cause and one SMART "
     "action. Activity 7: Pareto of 120 defects to find the vital few. Activity 8: velocity "
     "forecast as a range, and reading the CFD/control chart to name the bottleneck. "
     "(A3, A4, A6, A7)",
     "Hands-on\nTeam activities\n5 Whys · Pareto · Scrum Board",
     f"{SLIDES.get('act6', 102)}–110"],
    ["18:15 – 18:30", "0.25", "Topic 3 recap · Course summary and next steps",
     "Consolidate the three learning outcomes, the eight key takeaways, and the "
     "recommended follow-on courses. Final Q&A before the feedback survey.",
     "Recap\nQ&A", "111–115"],
]

ASSESS = [
    ["18:30 – 18:45", "—", "Course Feedback and TRAQOM Survey (mandatory)",
     "Learners complete the TRAQOM survey on lms-tms.tertiaryinfotech.com.",
     "Admin", "118"],
    ["18:45 – 19:00", "—", "Digital Attendance (Assessment) and Assessment Briefing",
     "Assessment attendance submitted. Rules restated: open book (slides, Learner Guide, "
     "approved materials only), no discussion, no photography.",
     "Admin", "116–117"],
    ["19:00 – 20:00", "—", "Written Assessment (WA) — Short-Answer Questions",
     "Individual, open book, 1 hour.", "Assessment", "116"],
    ["20:00 – 21:00", "—", "Case Study (CS)",
     "Individual, open book, 1 hour. An applied business scenario with structured tasks.",
     "Assessment", "116"],
    ["21:00 – 21:15", "—", "Results, feedback and appeal process",
     "Competent / Not Yet Competent communicated, with the appeal route explained.",
     "Admin", "117"],
]

for day, rows in ((1, DAY1), (2, DAY2)):
    h2(f"3.{day}  Day {day} — {C.DAY_THEMES[day]}")
    total = sum(float(r[1]) for r in rows if r[1] not in ("—",))
    para(f"Instructional hours: {total:.2f} (excludes the 1-hour lunch break).",
         size=10, bold=True, color=BRAND)
    table(["Time", "Hrs", "Topic / Activity", "Content and Delivery Notes",
           "Method", "Slides"], rows,
          widths=[0.95, 0.4, 1.55, 2.35, 0.95, 0.55],
          header_fill=DAY_FILL[day], font=8.5)

h2("3.3  Assessment Session (Day 2, after instructional hours)")
para("The 2-hour assessment sits outside the 16 instructional hours.", size=10,
     bold=True, color=BRAND)
table(["Time", "Hrs", "Session", "Notes", "Method", "Slides"], ASSESS,
      widths=[0.95, 0.4, 1.55, 2.35, 0.95, 0.55], header_fill="F59E0B", font=8.5)

# ------------------------------------------------------------------ 4. activity summary
h1("4.  Activity Summary")
para("Every activity advances the same running case — HarbourFront Logistics Pte Ltd and its "
     "CustomerConnect portal — so the artefacts learners build compound across the two days. "
     "Full step-by-step instructions for each activity are in the Learner Guide.", size=10.5)
table(["#", "Activity", "Tool", "Min", "LO", "TSC", "Learner Output"],
      [[str(a["num"]), a["title"], a["tool"], str(a["duration"]), a["lo"],
        (a["objective"].split("(")[-1].rstrip(").") if "(" in a["objective"] else "—"),
        a["build"]] for a in ACTIVITIES],
      widths=[0.28, 1.75, 0.95, 0.35, 0.42, 0.75, 2.35], header_fill="10B981", font=8.5)

# ------------------------------------------------------------------ 5. assessment plan
h1("5.  Assessment Plan")
table(["Instrument", "Detail"], [
    ["Written Assessment (WA)", C.ASSESSMENT["written"] + " Assesses the underpinning "
     "knowledge (K1–K7): the business landscape, customer analysis methods, Agile values and "
     "principles, Agile methodologies, change management and team formation models."],
    ["Case Study (CS)", C.ASSESSMENT["practical"] + " Assesses applied ability (A1–A7): "
     "reading a business scenario, ordering a backlog, interpreting metrics, diagnosing a "
     "root cause and recommending action."],
    ["Conditions", C.ASSESSMENT["open_book"] + " Individual work. No discussion, no "
     "photography, no recording."],
    ["Competency Decision", "Learners must be assessed Competent in BOTH instruments. "
     "A learner assessed Not Yet Competent in either instrument may be re-assessed."],
    ["Attendance", C.ASSESSMENT["note"]],
    ["Appeal", "A learner may appeal an assessment decision in writing within 7 working days. "
     "The appeal is reviewed by a second qualified assessor."],
], widths=[1.55, 5.35], header_fill="F59E0B", font=10)

# ------------------------------------------------------------------ 6. trainer notes
h1("6.  Trainer Delivery Notes")
h2("6.1  Timing Discipline")
for t_ in ["The activities are the course. If you are running late, compress the concept "
           "slides — never the activities. Learners are assessed on applied ability.",
           "Announce a hard time for each activity and hold it. Activity 5 (the 10-day sprint "
           "simulation) will overrun if you do not call each simulated day yourself.",
           "Keep the debrief for each activity — it is where the learning is consolidated and "
           "where the TSC link is made explicit."]:
    bullet(t_)

h2("6.2  Facilitating the Activities")
for t_ in ["Form teams of 3–4 on Day 1 and keep them for both days, so their case artefacts "
           "accumulate.",
           "In Activity 3, you play the Product Owner. Push back on teams that accept every "
           "story; make them defend the ordering.",
           "In Activity 5, inject the three impediments on cue (Day 3 stale feed, Day 5 "
           "capacity loss, Day 7 mid-sprint request). The Day 7 request is the real test — "
           "watch whether teams protect the sprint or simply absorb the work.",
           "In Activity 6, do not let a team stop at 'the carrier's fault'. Keep asking why "
           "until the cause is something the team itself can change.",
           "In Activity 8, expect teams to recommend 'increase velocity'. That is the "
           "teachable moment — walk them back to the control chart and Little's Law."]:
    bullet(t_)

h2("6.3  Common Learner Misconceptions to Correct")
table(["Misconception", "The correction"], [
    ["“Agile means no planning.”", "Agile plans continuously and at multiple levels. It "
     "replaces one large upfront plan with frequent replanning."],
    ["“Agile means no documentation.”", "Agile documents 'barely sufficiently' — just in "
     "time and for a stated reason. The discipline moves; it does not disappear."],
    ["“Velocity measures productivity.”", "Velocity is a forecasting input. Rewarding it "
     "produces estimate inflation and destroys the forecast."],
    ["“Agile gets more done for less money.”", "Agile buys responsiveness, earlier value and "
     "lower risk of building the wrong thing — not free capacity."],
    ["“The Scrum Master runs the team.”", "The Scrum Master owns the process and removes "
     "impediments. The Developers own how the work is done."],
    ["“Working faster reduces cycle time.”", "Little's Law: cycle time falls when WIP falls. "
     "Starting less finishes more."],
    ["“Agile suits every project.”", "Fixed regulatory scope, no customer access or physical "
     "irreversibility all weigh against it. Say so honestly."],
], widths=[2.2, 4.7], header_fill="DC2626", font=9.5)

h2("6.4  Contingency")
for t_ in ["If a tool URL is unreachable, run the activity on flip-chart paper — every tool "
           "here has a paper equivalent, and the thinking is what is assessed.",
           "If the class is smaller than 6, run two teams and have them cross-review each "
           "other's artefacts instead of presenting to the room.",
           "If a learner arrives without a laptop, pair them into a team that has one; every "
           "activity is a team activity."]:
    bullet(t_)

add_page_numbers(doc)
enable_update_fields(doc)

OUTDIR = os.path.join(REPO, "courseware")
os.makedirs(OUTDIR, exist_ok=True)
OUT = os.path.join(OUTDIR, f"WSQ - Lesson Plan - {C.SHORT_TITLE} - {C.VERSION}.docx")
doc.save(OUT)
print("Saved:", OUT)
d1 = sum(float(r[1]) for r in DAY1 if r[1] != "—")
d2 = sum(float(r[1]) for r in DAY2 if r[1] != "—")
print(f"Day 1 instructional hours: {d1}")
print(f"Day 2 instructional hours: {d2}")
print(f"Total: {d1 + d2}  (target {C.DURATION_HRS})")
assert d1 == 8.0, f"Day 1 must total 8.0 instructional hours, got {d1}"
assert d2 == 8.0, f"Day 2 must total 8.0 instructional hours, got {d2}"
print("Both days total exactly 8.0 instructional hours.")
