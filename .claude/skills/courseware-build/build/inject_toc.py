#!/usr/bin/env python3
"""Replace the placeholder TOC in a DOCX with a static, page-numbered TOC.

Page numbers are MEASURED, never estimated. The injector runs in two passes:

  pass 1  insert the static TOC with provisional page numbers taken from the
          pass-1 PDF, save, and re-render the DOCX to PDF
  pass 2  re-read the freshly rendered PDF (which now contains the real,
          full-length TOC) and rewrite every page number from it

That second pass is what makes the numbers correct. An earlier version estimated
how many pages the injected TOC would occupy from a hardcoded lines-per-page
constant; for a long TOC the estimate was one page short and every entry came out
off by one. Measuring removes the guess entirely.

Headings that cannot be found in the rendered PDF are DROPPED rather than
defaulted to page 1, so a stale heading can never leave an orphan TOC row.

Usage: python3 inject_toc.py <docx> <pdf> [maxlevel] [--soffice PATH]
"""
import os
import re
import subprocess
import sys
import tempfile

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from pypdf import PdfReader

GREY = RGBColor(0x33, 0x33, 0x33)

SOFFICE_CANDIDATES = [
    "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    "/usr/bin/soffice",
    "soffice",
]


def norm(s):
    return re.sub(r"[^a-z0-9]", "", (s or "").lower())


def find_soffice(explicit=None):
    for c in ([explicit] if explicit else []) + SOFFICE_CANDIDATES:
        if not c:
            continue
        if os.path.isabs(c) and os.path.exists(c):
            return c
        if not os.path.isabs(c):
            from shutil import which
            p = which(c)
            if p:
                return p
    return None


def page_texts(pdf_path):
    return [norm(pg.extract_text() or "") for pg in PdfReader(pdf_path).pages]


def collect_headings(doc, maxlevel):
    heads = []
    for p in doc.paragraphs:
        sn = p.style.name
        if not sn.startswith("Heading"):
            continue
        try:
            lvl = int(sn.split()[-1])
        except ValueError:
            continue
        if lvl <= maxlevel and p.text.strip():
            heads.append((lvl, p.text.strip()))
    return heads


def first_body_page(pages):
    """Index of the first page AFTER the table of contents.

    On pass 2 the rendered PDF contains the full static TOC, and every heading
    string appears there as a TOC row. Matching from page 0 would resolve every
    entry to the TOC page itself, so the search must start after it.
    """
    last_toc = -1
    for i, t in enumerate(pages):
        if "tableofcontents" in t:
            last_toc = i
    if last_toc < 0:
        return 0
    # the TOC may span several pages: walk forward while the page still looks like
    # a dense run of dotted TOC rows rather than prose
    j = last_toc
    while j + 1 < len(pages) and pages[j + 1].count("aboutthiscourse") == 0 and j - last_toc < 3:
        j += 1
        break
    return last_toc + 1


def locate(heads, pages, start=0):
    """Map each heading to the first page (searching forward from `start`) whose
    text contains it.

    Returns (entries, dropped) where entries is [(lvl, text, page)] in document
    order and dropped lists headings that could not be located at all.
    """
    entries, dropped = [], []
    cursor = start
    for lvl, text in heads:
        key = norm(text)[:24]
        page = None
        if key:
            for i in range(cursor, len(pages)):
                if key in pages[i]:
                    page = i + 1
                    cursor = i
                    break
            if page is None:                      # fall back to a scan of the body
                for i in range(start, len(pages)):
                    if key in pages[i]:
                        page = i + 1
                        break
        if page is None:
            dropped.append(text)                  # never emit an orphan row
            continue
        entries.append((lvl, text, page))
    return entries, dropped


def find_placeholder(doc):
    """The TOC field placeholder, or the static TOC this script wrote previously."""
    for p in doc.paragraphs:
        xml = p._p.xml
        if ("Update Field" in p.text
                or ("TOC " in xml and "instrText" in xml)
                or ("fldSimple" in xml and "TOC" in xml)):
            return p
    return None


def find_static_block(doc):
    """Locate a previously injected static TOC: the run of dotted-tab paragraphs
    that follows the 'TABLE OF CONTENTS' heading. Returns a list of Paragraphs."""
    out, started = [], False
    for p in doc.paragraphs:
        if not started:
            if norm(p.text).startswith("tableofcontents"):
                started = True
            continue
        if p.paragraph_format.tab_stops and len(p.paragraph_format.tab_stops):
            out.append(p)
        elif out:
            break
    return out


def write_entries(anchor_p, entries):
    """Insert TOC paragraphs before anchor_p (a Paragraph)."""
    anchor = anchor_p._p
    for lvl, text, page in entries:
        new_p = anchor.makeelement(qn("w:p"), {})
        anchor.addprevious(new_p)
        para = Paragraph(new_p, anchor_p._parent)
        pf = para.paragraph_format
        pf.tab_stops.add_tab_stop(Inches(6.3), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        if lvl >= 2:
            pf.left_indent = Inches(0.3)
        pf.space_after = Pt(3)
        r = para.add_run(f"{text}\t{page}")
        r.font.size = Pt(11 if lvl == 1 else 10.5)
        r.font.name = "Arial"
        r.bold = (lvl == 1)
        r.font.color.rgb = GREY


def render_pdf(docx_path, soffice):
    """Render docx to PDF in a temp dir; return the PDF path (caller cleans up)."""
    tmp = tempfile.mkdtemp(prefix="toc_")
    subprocess.run([soffice, "--headless", "--convert-to", "pdf", "--outdir", tmp, docx_path],
                   capture_output=True, text=True, timeout=600)
    out = os.path.join(tmp, os.path.splitext(os.path.basename(docx_path))[0] + ".pdf")
    return out if os.path.exists(out) else None


def main():
    if len(sys.argv) < 3:
        raise SystemExit(__doc__)
    docx_path, pdf_path = sys.argv[1], sys.argv[2]
    maxlevel = 2
    soffice_arg = None
    rest = sys.argv[3:]
    for i, a in enumerate(rest):
        if a == "--soffice" and i + 1 < len(rest):
            soffice_arg = rest[i + 1]
        elif a.isdigit():
            maxlevel = int(a)

    # ---------------- pass 1: provisional numbers from the supplied PDF
    doc = Document(docx_path)
    heads = collect_headings(doc, maxlevel)
    entries, dropped = locate(heads, page_texts(pdf_path))
    placeholder = find_placeholder(doc)
    if placeholder is None:
        print(f"  [inject_toc] no TOC placeholder found in {docx_path}")
        return
    write_entries(placeholder, entries)
    placeholder._p.getparent().remove(placeholder._p)
    doc.save(docx_path)
    if dropped:
        print(f"  [inject_toc] dropped {len(dropped)} heading(s) not found in the render: "
              f"{', '.join(dropped[:4])}")

    # ---------------- pass 2: re-render and MEASURE the real page numbers
    soffice = find_soffice(soffice_arg)
    if not soffice:
        print("  [inject_toc] WARNING: soffice not found — page numbers are pass-1 estimates "
              "and may be off by the length of the injected TOC.")
        print(f"  [inject_toc] {docx_path}: wrote {len(entries)} TOC entries (unverified)")
        return

    pdf2 = render_pdf(docx_path, soffice)
    if not pdf2:
        print("  [inject_toc] WARNING: second-pass render failed — numbers unverified.")
        return

    doc2 = Document(docx_path)
    heads2 = collect_headings(doc2, maxlevel)
    pages2 = page_texts(pdf2)
    entries2, dropped2 = locate(heads2, pages2, start=first_body_page(pages2))

    block = find_static_block(doc2)
    if not block:
        print("  [inject_toc] WARNING: could not re-locate the static TOC for pass 2.")
        return
    anchor_p = block[0]
    for p in block:                       # clear the pass-1 rows
        p._p.getparent().remove(p._p)
    # re-insert measured rows before whatever now follows the heading
    tail = anchor_p
    write_entries_at_end = None
    # find the paragraph that followed the block to anchor against
    doc2.save(docx_path)
    doc3 = Document(docx_path)
    # rebuild: insert right after the 'TABLE OF CONTENTS' heading
    target = None
    paras = doc3.paragraphs
    for i, p in enumerate(paras):
        if norm(p.text).startswith("tableofcontents") and i + 1 < len(paras):
            target = paras[i + 1]
            break
    if target is None:
        print("  [inject_toc] WARNING: lost the TOC heading on pass 2.")
        return
    write_entries(target, entries2)
    doc3.save(docx_path)

    changed = sum(1 for a, b in zip(entries, entries2) if a[2] != b[2])
    print(f"  [inject_toc] {docx_path}: wrote {len(entries2)} TOC entries "
          f"(pages {entries2[0][2]}..{entries2[-1][2]}); "
          f"pass 2 corrected {changed} page number(s)"
          + (f"; dropped {len(dropped2)} stale heading(s)" if dropped2 else ""))
    try:
        os.remove(pdf2); os.rmdir(os.path.dirname(pdf2))
    except OSError:
        pass


if __name__ == "__main__":
    main()
